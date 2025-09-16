import dash
from dash import html, Input, Output, State, dcc, callback_context
import dash_bootstrap_components as dbc
from model_serving_utils import query_endpoint
import json
import os
import time
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from databricks.sdk import WorkspaceClient
import threading
import io
from docx import Document
# Removed dbldatagen import - now using Databricks job approach

class SyntheticDataGenerator:
    def __init__(self, app, endpoint_name):
        self.app = app
        self.endpoint_name = endpoint_name
        self.volume_path = "/Volumes/conor_smith/synthetic_data_app/synthetic_data_volume"
        
        # Progress tracking
        self.generation_state = {
            'active': False,
            'current_doc': 0,
            'total_docs': 0,
            'current_step': '',
            'completed_files': [],
            'error': None,
            'start_time': None,
            'estimated_total_time': 0
        }
        
        self._create_callbacks()
        self._add_custom_css()

    def _estimate_total_time(self, operations):
        """Calculate estimated total time for all operations based on type and configuration."""
        total_time = 0
        
        for operation in operations:
            op_type = operation.get('type', '')
            config = operation.get('config', {})
            
            if op_type == 'pdf':
                # PDF operations
                count = config.get('count', 1)
                include_images = config.get('include_images', False)
                if include_images:
                    total_time += count * 35  # 35 seconds per PDF with images
                else:
                    total_time += count * 13  # 13 seconds per PDF without images
                    
            elif op_type == 'text':
                # Text operations (treat similar to PDF without images)
                count = config.get('count', 1)
                total_time += count * 13  # 13 seconds per text document
                
            elif op_type == 'tabular':
                # Tabular operations
                row_count = config.get('row_count', 1000)
                columns = config.get('columns', [])
                
                # Check if any column is GenAI Text
                has_genai = any(col.get('data_type') == 'GenAI Text' for col in columns)
                
                if has_genai:
                    if row_count <= 500:
                        total_time += 22  # 22 seconds for GenAI ≤ 500 rows
                    else:
                        total_time += 52  # 52 seconds for GenAI > 500 rows
                else:
                    total_time += 21  # 21 seconds for tabular without GenAI
        
        return total_time

    def _parse_column_references(self, prompt_text):
        """Extract column names referenced in the prompt using <Column Name> syntax."""
        import re
        # Find all text within angle brackets
        column_refs = re.findall(r'<([^<>]+)>', prompt_text)
        return column_refs
    
    # Removed _substitute_column_references_spark - now handled by Databricks job
    
    def _substitute_column_references_pandas(self, prompt_template, row_data, columns):
        """Substitute column references in prompt for a single row."""
        # Parse column references from prompt
        column_refs = self._parse_column_references(prompt_template)
        
        if not column_refs:
            # No column references, return original prompt
            return prompt_template
        
        # Build list of available non-GenAI columns
        available_columns = {}
        for col in columns:
            col_name = col.get('name', '')
            col_type = col.get('data_type', '')
            if col_type != 'GenAI Text' and col_name:
                available_columns[col_name] = col_name
        
        # Start with the original prompt
        substituted_prompt = prompt_template
        
        # Replace each column reference with the actual value
        for col_ref in column_refs:
            if col_ref in available_columns and col_ref in row_data:
                placeholder = f"<{col_ref}>"
                # Get the value and convert to string
                value = str(row_data[col_ref]) if row_data[col_ref] is not None else ""
                substituted_prompt = substituted_prompt.replace(placeholder, value)
            else:
                if col_ref not in available_columns:
                    print(f"Warning: Column '{col_ref}' not found or is a GenAI Text column. Leaving placeholder unchanged.")
                elif col_ref not in row_data:
                    print(f"Warning: Column '{col_ref}' not found in row data. Leaving placeholder unchanged.")
        
        return substituted_prompt

    def _create_callbacks(self):
        # Enable/disable add operation button based on data type selection
        @self.app.callback(
            Output('add-operation-button', 'disabled'),
            Input('data-type-selector', 'value'),
            prevent_initial_call=True
        )
        def enable_add_operation(data_type):
            return data_type is None
        
        # Add operation callback
        @self.app.callback(
            Output('operations-store', 'data'),
            Output('data-type-selector', 'value'),
            Input('add-operation-button', 'n_clicks'),
            State('data-type-selector', 'value'),
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def add_operation(n_clicks, data_type, operations):
            if not n_clicks or not data_type:
                return dash.no_update, dash.no_update
            
            operations = operations or []
            
            # Create new operation with unique ID
            operation_id = f"op_{len(operations)}_{int(time.time())}"
            new_operation = {
                'id': operation_id,
                'type': data_type,
                'configured': False,
                'config': {}
            }
            
            operations.append(new_operation)
            
            # Clear the selector
            return operations, None
        
        # Display operations callback
        @self.app.callback(
            Output('operations-container', 'children'),
            Output('generate-all-button', 'disabled'),
            Input('operations-store', 'data'),
            prevent_initial_call=True
        )
        def display_operations(operations):
            if not operations:
                return [], True
            
            operation_cards = []
            configured_count = 0
            
            for op in operations:
                card = self._create_operation_card(op)
                operation_cards.append(card)
                if op.get('configured', False):
                    configured_count += 1
            
            # Enable generate button if at least one operation is configured
            generate_disabled = configured_count == 0
            
            return operation_cards, generate_disabled

        # Start batch generation callback
        @self.app.callback(
            Output('progress-interval', 'disabled'),
            Output('progress-store', 'data'),
            Output('generate-all-button', 'disabled', allow_duplicate=True),
            Input('generate-all-button', 'n_clicks'),
            State('operations-store', 'data'),
            State('company-name', 'value'),
            State('company-sector', 'value'),
            prevent_initial_call=True
        )
        def start_batch_generation(n_clicks, operations, company_name, company_sector):
            if not n_clicks or not operations:
                return dash.no_update, dash.no_update, dash.no_update
            
            # Small delay to ensure any pending column type updates have completed
            import time
            time.sleep(0.1)  # 100ms delay to allow callback queue to process
            
            # Debug: Log column types before generation
            print("🔍 DEBUG: Operations before generation:")
            for i, op in enumerate(operations):
                if op.get('type') == 'tabular' and 'columns' in op.get('config', {}):
                    print(f"   Operation {i+1}: {op.get('config', {}).get('table_name', 'unnamed')}")
                    for col in op['config']['columns']:
                        print(f"      - {col.get('name', 'unnamed')}: {col.get('data_type', 'unknown')}")
            
            # Filter only configured operations
            configured_ops = [op for op in operations if op.get('configured', False)]
            if not configured_ops:
                return dash.no_update, dash.no_update, dash.no_update
            
            # Calculate estimated total time for countdown
            estimated_time = self._estimate_total_time(configured_ops)
            
            # Reset and start generation state
            self.generation_state = {
                'active': True,
                'current_operation': 0,
                'total_operations': len(configured_ops),
                'current_step': 'Initializing batch...',
                'completed_items': [],
                'error': None,
                'start_time': time.time(),
                'estimated_total_time': estimated_time,
                'operations': configured_ops,
                'company_name': company_name or 'Acme Solutions Inc.',
                'company_sector': company_sector or 'technology'
            }
            
            # Start generation in background thread
            threading.Thread(
                target=self._generate_batch_background, 
                args=(configured_ops, company_name, company_sector),
                daemon=True
            ).start()
            
            # Enable interval and disable button
            return False, {'status': 'started'}, True
        
        # Progress update callback
        @self.app.callback(
            Output('current-generation-status', 'children'),
            Output('progress-interval', 'disabled', allow_duplicate=True),
            Output('generate-all-button', 'disabled', allow_duplicate=True),
            Output('generation-store', 'data'),
            Input('progress-interval', 'n_intervals'),
            prevent_initial_call=True
        )
        def update_progress(n_intervals):
            if not self.generation_state.get('active', False) and self.generation_state.get('current_operation', 0) == 0:
                return dash.no_update, dash.no_update, dash.no_update, dash.no_update
            
            # Calculate time remaining for countdown
            start_time = self.generation_state.get('start_time', time.time())
            estimated_time = self.generation_state.get('estimated_total_time', 0)
            elapsed_time = time.time() - start_time
            time_remaining = max(0, estimated_time - elapsed_time)
            
            # Check if generation is complete
            if not self.generation_state.get('active', False):
                if self.generation_state.get('error'):
                    # Error state
                    error_message = dbc.Alert([
                        html.H5("❌ Batch Generation Failed", className="mb-2"),
                        html.P(f"Error: {self.generation_state['error']}")
                    ], color="danger")
                    return error_message, True, False, {'status': 'error', 'error': self.generation_state['error']}
                else:
                    # Success state - batch completed
                    elapsed_time = time.time() - self.generation_state['start_time']
                    total_items = len(self.generation_state.get('completed_items', []))
                    
                    success_message = dbc.Alert([
                        dbc.Row([
                            dbc.Col([
                                html.H5("✅ Batch Generation Complete!", className="mb-2"),
                                html.P(f"Successfully generated {total_items} item(s) from {self.generation_state.get('total_operations', 0)} operation(s) in {elapsed_time:.1f} seconds"),
                                html.P("You can add more operations above or generate the current batch again!", className="mb-0 text-muted")
                            ], width=8),
                            dbc.Col([
                                dbc.Button([
                                    html.I(className="fas fa-download me-2"),
                                    "Download All Files"
                                ], 
                                id="download-all-files-btn",
                                color="primary", 
                                size="lg",
                                className="w-100",
                                disabled=False,
                                title="Download all files generated in this batch as a ZIP file (downloads from volume to local directory first, then creates ZIP)")
                            ], width=4)
                        ])
                    ], color="success")
                    
                    return success_message, True, False, {'status': 'complete', 'items': self.generation_state.get('completed_items', [])}
            
            # Active generation state
            current_op = self.generation_state.get('current_operation', 0)
            total_ops = self.generation_state.get('total_operations', 1)
            
            status_message = dbc.Alert([
                html.Div([
                    html.Div([
                        html.I(className="fas fa-spinner fa-spin me-2"),
                        html.H5(f"🔄 Processing Batch Operations...", className="d-inline mb-0")
                    ], className="d-flex align-items-center mb-3"),
                    
                    html.P(f"Operation {current_op} of {total_ops}", className="mb-2"),
                    html.P(f"Current step: {self.generation_state.get('current_step', 'Processing...')}", className="mb-3 text-muted"),
                    
                    # Countdown timer display
                    html.Div([
                        html.H3(f"{int(time_remaining // 60)}:{int(time_remaining % 60):02d}", 
                               className="countdown-timer text-center mb-1"),
                        html.P("Estimated Time Remaining", className="text-center mb-0 small text-muted")
                    ], className="mb-2")
                ])
            ], color="info")
            
            return status_message, False, True, {'status': 'generating', 'time_remaining': time_remaining, 'elapsed_time': elapsed_time}
        
        # History update callback
        @self.app.callback(
            Output('generation-history', 'children'),
            Output('history-store', 'data'),
            Input('generation-store', 'data'),
            State('history-store', 'data'),
            prevent_initial_call=True
        )
        def update_history(generation_data, history_data):
            if not generation_data or generation_data.get('status') != 'complete':
                return dash.no_update, dash.no_update
            
            # Add new generation to history
            history_data = history_data or []
            
            new_entry = {
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'data_type': self.generation_state.get('data_type', 'unknown'),
                'description': self.generation_state.get('description', 'No description'),
                'company_name': self.generation_state.get('company_name', 'Unknown Company'),
                'company_sector': self.generation_state.get('company_sector', 'unknown'),
                'items': self.generation_state.get('completed_items', [])
            }
            
            history_data.append(new_entry)
            
            # Build history display
            history_cards = []
            for i, entry in enumerate(reversed(history_data[-10:])):  # Show last 10 entries, most recent first
                history_cards.append(
                    dbc.Card([
                        dbc.CardBody([
                            html.Div([
                                html.H6(f"{self._format_data_type(entry['data_type'])}", className="card-title mb-1"),
                                html.P(f"Company: {entry['company_name']} ({entry['company_sector']})", className="mb-1 small text-muted"),
                                html.P(f"Description: {entry['description']}", className="mb-1"),
                                html.P(f"Generated: {entry['timestamp']}", className="mb-1 small text-muted"),
                                self._format_items_display(entry['items'], entry['data_type'])
                            ])
                        ])
                    ], className="mb-2")
                )
            
            if history_cards:
                history_display = [
                    html.H4("Generation History", className="mb-3"),
                    html.Div(history_cards)
                ]
            else:
                history_display = []
            
            return history_display, history_data
        
        # Dynamic callback for operation configuration updates
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            [Input({'type': 'pdf-doc-name', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'pdf-description', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'pdf-doc-type', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'pdf-count', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'pdf-include-images', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'text-doc-name', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'text-description', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'text-doc-type', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'text-format', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'text-count', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'tabular-name', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'tabular-rows', 'index': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-name', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-type', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-min', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-max', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-min-date', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'date'),
             Input({'type': 'col-max-date', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'date'),
             Input({'type': 'col-prompt', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'custom-value-input', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL, 'idx': dash.dependencies.ALL}, 'value'),
             Input({'type': 'custom-weight-input', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL, 'idx': dash.dependencies.ALL}, 'value'),
             Input({'type': 'use-weights-checkbox', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-ordered', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-country-code', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
             Input({'type': 'col-coords-only', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value')],
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def update_operation_configs(*args):
            operations = args[-1]  # Last argument is the operations state
            if not operations:
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update
            
            # Parse the triggered input
            triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
            if triggered_id == '':
                return dash.no_update
            
            try:
                import json
                triggered_comp = json.loads(triggered_id)
                comp_type = triggered_comp['type']
                new_value = ctx.triggered[0]['value']
                
                # Handle column updates differently from operation updates
                if comp_type in ['col-name', 'col-type', 'col-min', 'col-max', 'col-min-date', 'col-max-date', 
                                'col-prompt', 'col-max-tokens', 'custom-value-input', 'custom-weight-input', 'use-weights-checkbox', 'col-ordered',
                                'col-country-code', 'col-coords-only']:
                    op_id = triggered_comp['op']
                    col_id = triggered_comp['col']
                    
                    # Find and update the column
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            if 'columns' not in op['config']:
                                op['config']['columns'] = []
                            
                            # Find or create the column
                            column = None
                            for col in op['config']['columns']:
                                if col['id'] == col_id:
                                    column = col
                                    break
                            
                            if column:
                                # Update the column property
                                if comp_type == 'col-name':
                                    column['name'] = new_value
                                elif comp_type == 'col-type':
                                    print(f"🔄 Column type changed: {column.get('name', 'unnamed')} → {new_value}")
                                    column['data_type'] = new_value
                                    # Initialize type-specific fields when type changes
                                    if new_value == 'Integer':
                                        if 'min_value' not in column:
                                            column['min_value'] = 1
                                        if 'max_value' not in column:
                                            column['max_value'] = 100
                                        if 'ordered_values' not in column:
                                            column['ordered_values'] = False
                                    elif new_value == 'Date':
                                        if 'min_date' not in column:
                                            column['min_date'] = '2020-01-01'
                                        if 'max_date' not in column:
                                            column['max_date'] = '2024-12-31'
                                        if 'ordered_values' not in column:
                                            column['ordered_values'] = False
                                    elif new_value == 'GenAI Text':
                                        if 'prompt' not in column:
                                            column['prompt'] = ''
                                        if 'max_tokens' not in column:
                                            column['max_tokens'] = 500
                                    elif new_value == 'Custom Values':
                                        if 'custom_values' not in column:
                                            column['custom_values'] = ['']
                                        if 'use_weights' not in column:
                                            column['use_weights'] = False
                                        if 'custom_weights' not in column:
                                            column['custom_weights'] = [1]
                                        if 'ordered_values' not in column:
                                            column['ordered_values'] = False
                                    elif new_value == 'Country Defined Coordinates':
                                        if 'country_code' not in column:
                                            column['country_code'] = 'US'
                                        if 'coords_only' not in column:
                                            column['coords_only'] = False
                                elif comp_type == 'col-min':
                                    column['min_value'] = new_value
                                elif comp_type == 'col-max':
                                    column['max_value'] = new_value
                                elif comp_type == 'col-min-date':
                                    column['min_date'] = new_value
                                elif comp_type == 'col-max-date':
                                    column['max_date'] = new_value
                                elif comp_type == 'col-prompt':
                                    column['prompt'] = new_value
                                elif comp_type == 'col-max-tokens':
                                    column['max_tokens'] = new_value
                                elif comp_type == 'custom-value-input':
                                    # Handle custom value input updates
                                    idx = triggered_comp['idx']
                                    custom_values = column.get('custom_values', [''])
                                    # Ensure list is long enough
                                    while len(custom_values) <= idx:
                                        custom_values.append('')
                                    custom_values[idx] = new_value or ''
                                    column['custom_values'] = custom_values
                                elif comp_type == 'custom-weight-input':
                                    # Handle custom weight input updates
                                    idx = triggered_comp['idx']
                                    custom_weights = column.get('custom_weights', [1])
                                    # Ensure list is long enough
                                    while len(custom_weights) <= idx:
                                        custom_weights.append(1)
                                    try:
                                        custom_weights[idx] = float(new_value) if new_value else 1.0
                                    except (ValueError, TypeError):
                                        custom_weights[idx] = 1.0
                                    column['custom_weights'] = custom_weights
                                elif comp_type == 'use-weights-checkbox':
                                    # Handle weights checkbox toggle
                                    use_weights = bool(new_value and 'use_weights' in new_value)
                                    column['use_weights'] = use_weights
                                elif comp_type == 'col-ordered':
                                    # Handle ordered values checkbox toggle
                                    ordered_values = bool(new_value and 'ordered' in new_value)
                                    column['ordered_values'] = ordered_values
                                elif comp_type == 'col-country-code':
                                    # Handle country code selection for Country Defined Coordinates
                                    column['country_code'] = new_value if new_value else 'US'
                                elif comp_type == 'col-coords-only':
                                    # Handle coordinates only checkbox toggle
                                    coords_only = bool(new_value and 'coords_only' in new_value)
                                    column['coords_only'] = coords_only
                                
                                # After updating any column, check if we need to clamp row count due to GenAI presence
                                if comp_type == 'col-type':
                                    print(f"🔍 Column type changed, checking row count limits...")
                                    current_row_count = op['config'].get('row_count', 1000)
                                    has_genai = any(col.get('data_type') == 'GenAI Text' for col in op['config'].get('columns', []))
                                    max_allowed = 1000 if has_genai else 1000000
                                    
                                    print(f"   - Current row count: {current_row_count}")
                                    print(f"   - Has GenAI after change: {has_genai}")
                                    print(f"   - Max allowed: {max_allowed}")
                                    
                                    if current_row_count > max_allowed:
                                        print(f"   - 🔧 AUTO-CLAMPING: {current_row_count:,} → {max_allowed:,} due to GenAI Text column")
                                        op['config']['row_count'] = max_allowed
                                    else:
                                        print(f"   - ✅ Row count {current_row_count:,} is within limits")
                                
                                # Update configured status
                                table_name = op['config'].get('table_name', '')
                                op['configured'] = bool(table_name and table_name.strip() and len(op['config']['columns']) > 0)
                            break
                else:
                    # Handle operation-level updates
                    op_id = triggered_comp['index']
                    
                    # Find and update the operation
                    for op in operations:
                        if op['id'] == op_id:
                            if 'config' not in op:
                                op['config'] = {}
                            
                            # Update the specific config value
                            if comp_type == 'pdf-doc-name':
                                op['config']['doc_name'] = new_value
                            elif comp_type == 'pdf-description':
                                op['config']['description'] = new_value
                            elif comp_type == 'pdf-doc-type':
                                op['config']['doc_type'] = new_value
                            elif comp_type == 'pdf-count':
                                op['config']['count'] = new_value
                            elif comp_type == 'pdf-include-images':
                                op['config']['include_images'] = 'include_images' in (new_value or [])
                            elif comp_type == 'text-doc-name':
                                op['config']['doc_name'] = new_value
                            elif comp_type == 'text-description':
                                op['config']['description'] = new_value
                            elif comp_type == 'text-doc-type':
                                op['config']['doc_type'] = new_value
                            elif comp_type == 'text-format':
                                op['config']['file_format'] = new_value
                            elif comp_type == 'text-count':
                                op['config']['count'] = new_value
                            elif comp_type == 'tabular-name':
                                op['config']['table_name'] = new_value
                            elif comp_type == 'tabular-rows':
                                # Validate and clamp row count based on GenAI Text presence
                                print(f"🔧 MAIN CALLBACK: Processing row count change")
                                print(f"   - Raw new_value: {repr(new_value)} (type: {type(new_value)})")
                                print(f"   - Current stored value: {op['config'].get('row_count', 'NOT_SET')}")
                                
                                if new_value is not None and str(new_value).strip() != '':
                                    try:
                                        # Convert to integer if it's a string
                                        if isinstance(new_value, str):
                                            print(f"   - Converting string '{new_value}' to int")
                                            new_value = int(new_value)
                                        
                                        print(f"   - Processed value: {new_value}")
                                        
                                        # Check if operation has GenAI Text columns
                                        has_genai = any(col.get('data_type') == 'GenAI Text' for col in op['config'].get('columns', []))
                                        max_allowed = 1000 if has_genai else 1000000
                                        
                                        print(f"   - Has GenAI: {has_genai}")
                                        print(f"   - Max allowed: {max_allowed:,}")
                                        print(f"   - Input value: {new_value}")
                                        
                                        # Clamp the value
                                        if new_value < 1:
                                            clamped_value = 1
                                            print(f"   - Clamped {new_value} → {clamped_value} (minimum)")
                                        elif new_value > max_allowed:
                                            clamped_value = max_allowed
                                            limit_reason = "GenAI Text columns" if has_genai else "system limits"
                                            print(f"   - Clamped {new_value:,} → {clamped_value:,} due to {limit_reason}")
                                        else:
                                            clamped_value = new_value
                                            print(f"   - Value {new_value:,} is within limits")
                                        
                                        print(f"   - Setting row_count to: {clamped_value}")
                                        op['config']['row_count'] = clamped_value
                                        
                                    except (ValueError, TypeError) as e:
                                        print(f"   - ERROR converting '{new_value}': {e}")
                                        print(f"   - Keeping current: {op['config'].get('row_count', 1000)}")
                                        # Don't change the stored value for invalid inputs
                                else:
                                    # For None or empty values during typing, don't change stored value
                                    current_stored = op['config'].get('row_count', 1000)
                                    print(f"   - None/empty value received, keeping current: {current_stored}")
                                    # Don't update the stored value
                            
                            # Mark as configured based on operation type
                            if op['type'] == 'tabular':
                                # For tabular operations, require table name and at least one column
                                table_name = op['config'].get('table_name', '')
                                columns = op['config'].get('columns', [])
                                op['configured'] = bool(table_name and table_name.strip() and len(columns) > 0)
                            elif op['type'] in ['pdf', 'text']:
                                # For PDF and text operations, require document name and description
                                doc_name = op['config'].get('doc_name', '')
                                description = op['config'].get('description', '')
                                op['configured'] = bool(doc_name and doc_name.strip() and description and description.strip())
                            else:
                                # For other operations, require description only
                                description = op['config'].get('description', '')
                                op['configured'] = bool(description and description.strip())
                            break
                
                return operations
                
            except Exception as e:
                print(f"Error updating operation config: {str(e)}")
                return dash.no_update
        
        # Remove operation callback
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            Input({'type': 'remove-op', 'index': dash.dependencies.ALL}, 'n_clicks'),
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def remove_operation(n_clicks_list, operations):
            if not operations or not any(n_clicks_list):
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update
            
            # Find which remove button was clicked
            triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
            if triggered_id == '':
                return dash.no_update
            
            try:
                import json
                triggered_comp = json.loads(triggered_id)
                op_id_to_remove = triggered_comp['index']
                
                # Remove the operation with matching ID
                operations = [op for op in operations if op['id'] != op_id_to_remove]
                
                return operations
                
            except Exception as e:
                print(f"Error removing operation: {str(e)}")
                return dash.no_update
        
        # Add column callback
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            Input({'type': 'add-column', 'index': dash.dependencies.ALL}, 'n_clicks'),
            State('operations-store', 'data'),
            State({'type': 'col-name', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-type', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-min', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-max', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-prompt', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-max-tokens', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            prevent_initial_call=True
        )
        def add_column(n_clicks_list, operations, col_names, col_types, col_mins, col_maxs, col_prompts, col_max_tokens):
            if not operations or not any(n_clicks_list):
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update
            
            try:
                import json
                triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
                triggered_comp = json.loads(triggered_id)
                op_id = triggered_comp['index']
                
                # First, update existing columns with current input values to preserve debounced inputs
                # This captures any values that were typed but not yet saved due to debouncing
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        columns = op['config'].get('columns', [])
                        
                        # Update each existing column with current input state values
                        for i, col in enumerate(columns):
                            # Update name if we have a corresponding value
                            if col_names and i < len(col_names) and col_names[i] is not None:
                                col['name'] = col_names[i]
                            
                            # Update data type if we have a corresponding value  
                            if col_types and i < len(col_types) and col_types[i] is not None:
                                new_type = col_types[i]
                                col['data_type'] = new_type
                                # Initialize type-specific fields when type changes
                                if new_type == 'Integer':
                                    if 'min_value' not in col:
                                        col['min_value'] = 1
                                    if 'max_value' not in col:
                                        col['max_value'] = 100
                                    if 'ordered_values' not in col:
                                        col['ordered_values'] = False
                                elif new_type == 'Date':
                                    if 'min_date' not in col:
                                        col['min_date'] = '2020-01-01'
                                    if 'max_date' not in col:
                                        col['max_date'] = '2024-12-31'
                                    if 'ordered_values' not in col:
                                        col['ordered_values'] = False
                                elif new_type == 'GenAI Text':
                                    if 'prompt' not in col:
                                        col['prompt'] = ''
                                    if 'max_tokens' not in col:
                                        col['max_tokens'] = 1000
                            
                            # Update min/max values for Integer types
                            if col.get('data_type') == 'Integer':
                                if col_mins and i < len(col_mins) and col_mins[i] is not None:
                                    col['min_value'] = col_mins[i]
                                if col_maxs and i < len(col_maxs) and col_maxs[i] is not None:
                                    col['max_value'] = col_maxs[i]
                            
                            # Update prompt and max_tokens for GenAI Text types
                            if col.get('data_type') == 'GenAI Text':
                                if col_prompts and i < len(col_prompts) and col_prompts[i] is not None:
                                    col['prompt'] = col_prompts[i]
                                if col_max_tokens and i < len(col_max_tokens) and col_max_tokens[i] is not None:
                                    col['max_tokens'] = col_max_tokens[i]
                        break
                
                # Now find the operation and add a new column
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        if 'columns' not in op['config']:
                            op['config']['columns'] = []
                        
                        # Create new column with unique ID
                        new_column = {
                            'id': f"col_{len(op['config']['columns'])}_{int(time.time())}",
                            'name': f"column_{len(op['config']['columns']) + 1}",
                            'data_type': 'Integer',
                            'min_value': 1,
                            'max_value': 100,
                            'min_date': '2020-01-01',
                            'max_date': '2024-12-31',
                            'prompt': '',
                            'max_tokens': 500,
                            'custom_values': [''],
                            'use_weights': False,
                            'custom_weights': [1],
                            'ordered_values': False
                        }
                        op['config']['columns'].append(new_column)
                        
                        # Update configured status
                        table_name = op['config'].get('table_name', '')
                        op['configured'] = bool(table_name and table_name.strip() and len(op['config']['columns']) > 0)
                        break
                
                return operations
                
            except Exception as e:
                print(f"Error adding column: {str(e)}")
                return dash.no_update
        
        # Remove column callback
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            Input({'type': 'remove-column', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'n_clicks'),
            State('operations-store', 'data'),
            State({'type': 'col-name', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-type', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-min', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-max', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-prompt', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            State({'type': 'col-max-tokens', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'value'),
            prevent_initial_call=True
        )
        def remove_column(n_clicks_list, operations, col_names, col_types, col_mins, col_maxs, col_prompts, col_max_tokens):
            if not operations or not any(n_clicks_list):
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update
            
            try:
                import json
                triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
                triggered_comp = json.loads(triggered_id)
                op_id = triggered_comp['op']
                col_id = triggered_comp['col']
                
                # First, preserve current input values before removing column
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        columns = op['config'].get('columns', [])
                        
                        # Update each existing column with current input state values
                        for i, col in enumerate(columns):
                            # Update name if we have a corresponding value
                            if col_names and i < len(col_names) and col_names[i] is not None:
                                col['name'] = col_names[i]
                            
                            # Update data type if we have a corresponding value  
                            if col_types and i < len(col_types) and col_types[i] is not None:
                                new_type = col_types[i]
                                col['data_type'] = new_type
                                # Initialize type-specific fields when type changes
                                if new_type == 'Integer':
                                    if 'min_value' not in col:
                                        col['min_value'] = 1
                                    if 'max_value' not in col:
                                        col['max_value'] = 100
                                    if 'ordered_values' not in col:
                                        col['ordered_values'] = False
                                elif new_type == 'Date':
                                    if 'min_date' not in col:
                                        col['min_date'] = '2020-01-01'
                                    if 'max_date' not in col:
                                        col['max_date'] = '2024-12-31'
                                    if 'ordered_values' not in col:
                                        col['ordered_values'] = False
                                elif new_type == 'GenAI Text':
                                    if 'prompt' not in col:
                                        col['prompt'] = ''
                                    if 'max_tokens' not in col:
                                        col['max_tokens'] = 1000
                            
                            # Update min/max values for Integer types
                            if col.get('data_type') == 'Integer':
                                if col_mins and i < len(col_mins) and col_mins[i] is not None:
                                    col['min_value'] = col_mins[i]
                                if col_maxs and i < len(col_maxs) and col_maxs[i] is not None:
                                    col['max_value'] = col_maxs[i]
                            
                            # Update prompt and max_tokens for GenAI Text types
                            if col.get('data_type') == 'GenAI Text':
                                if col_prompts and i < len(col_prompts) and col_prompts[i] is not None:
                                    col['prompt'] = col_prompts[i]
                                if col_max_tokens and i < len(col_max_tokens) and col_max_tokens[i] is not None:
                                    col['max_tokens'] = col_max_tokens[i]
                        break
                
                # Now find the operation and remove the column
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        if 'columns' in op['config']:
                            op['config']['columns'] = [col for col in op['config']['columns'] if col['id'] != col_id]
                            
                            # Update configured status
                            table_name = op['config'].get('table_name', '')
                            op['configured'] = bool(table_name and table_name.strip() and len(op['config']['columns']) > 0)
                        break
                
                return operations
                
            except Exception as e:
                print(f"Error removing column: {str(e)}")
                return dash.no_update
        
        # Update columns container callback
        @self.app.callback(
            Output({'type': 'columns-container', 'index': dash.dependencies.MATCH}, 'children'),
            Input('operations-store', 'data'),
            State({'type': 'columns-container', 'index': dash.dependencies.MATCH}, 'id'),
            prevent_initial_call=True
        )
        def update_columns_container(operations, container_id):
            if not operations:
                return []
            
            op_id = container_id['index']
            
            # Find the matching operation
            for op in operations:
                if op['id'] == op_id and op['type'] == 'tabular':
                    columns = op['config'].get('columns', [])
                    return self._create_column_cards(columns, op_id)
            
            return []
        
        # Update column type-specific configuration callback
        @self.app.callback(
            Output({'type': 'col-config', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'children'),
            Input({'type': 'col-type', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'value'),
            State({'type': 'col-config', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'id'),
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def update_column_config(col_type, config_id, operations):
            if col_type == 'Integer':
                # Get existing values for this column
                min_val, max_val = 1, 100
                
                if operations:
                    op_id = config_id['op']
                    col_id = config_id['col']
                    
                    # Find the operation and column to get existing values
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            columns = op['config'].get('columns', [])
                            for col in columns:
                                if col['id'] == col_id:
                                    min_val = col.get('min_value', 1)
                                    max_val = col.get('max_value', 100)
                                    break
                            break
                
                # Show min/max inputs for Integer type
                return [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Min Value:", className="form-label"),
                            dbc.Input(
                                id={'type': 'col-min', 'op': config_id['op'], 'col': config_id['col']},
                                type="number",
                                value=min_val,
                                size="sm",
                                debounce=True
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Max Value:", className="form-label"),
                            dbc.Input(
                                id={'type': 'col-max', 'op': config_id['op'], 'col': config_id['col']},
                                type="number",
                                value=max_val,
                                size="sm",
                                debounce=True
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dbc.Checklist(
                                id={'type': 'col-ordered', 'op': config_id['op'], 'col': config_id['col']},
                                options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                                value=['ordered'] if (operations and any(
                                    col.get('ordered_values', False) 
                                    for op in operations 
                                    if op.get('id') == config_id['op'] and op.get('type') == 'tabular'
                                    for col in op.get('config', {}).get('columns', [])
                                    if col.get('id') == config_id['col']
                                )) else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            ),
                            html.Small("When unchecked, values are randomized (default). When checked, values are generated in sequence.", 
                                      className="text-muted d-block", style={'fontSize': '12px'})
                        ], width=12)
                    ])
                ]
            elif col_type == 'Date':
                # Get existing date values for this column
                min_date, max_date = '2020-01-01', '2024-12-31'
                
                if operations:
                    op_id = config_id['op']
                    col_id = config_id['col']
                    
                    # Find the operation and column to get existing values
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            columns = op['config'].get('columns', [])
                            for col in columns:
                                if col['id'] == col_id:
                                    min_date = col.get('min_date', '2020-01-01')
                                    max_date = col.get('max_date', '2024-12-31')
                                    break
                            break
                
                # Show date pickers for Date type
                return [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Min Date:", className="form-label"),
                            dcc.DatePickerSingle(
                                id={'type': 'col-min-date', 'op': config_id['op'], 'col': config_id['col']},
                                date=min_date,
                                display_format='YYYY-MM-DD',
                                style={'width': '100%', 'fontSize': '14px'}
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Max Date:", className="form-label"),
                            dcc.DatePickerSingle(
                                id={'type': 'col-max-date', 'op': config_id['op'], 'col': config_id['col']},
                                date=max_date,
                                display_format='YYYY-MM-DD',
                                style={'width': '100%', 'fontSize': '14px'}
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dbc.Checklist(
                                id={'type': 'col-ordered', 'op': config_id['op'], 'col': config_id['col']},
                                options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                                value=['ordered'] if (operations and any(
                                    col.get('ordered_values', False) 
                                    for op in operations 
                                    if op.get('id') == config_id['op'] and op.get('type') == 'tabular'
                                    for col in op.get('config', {}).get('columns', [])
                                    if col.get('id') == config_id['col']
                                )) else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            ),
                            html.Small("When unchecked, dates are randomized (default). When checked, dates are generated in sequence.", 
                                      className="text-muted d-block", style={'fontSize': '12px'})
                        ], width=12)
                    ])
                ]
            elif col_type == 'GenAI Text':
                # Get existing prompt value for this column
                prompt_val = ""
                
                if operations:
                    op_id = config_id['op']
                    col_id = config_id['col']
                    
                    # Find the operation and column to get existing prompt and max_tokens
                    max_tokens_val = 500  # Default value
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            columns = op['config'].get('columns', [])
                            for col in columns:
                                if col['id'] == col_id:
                                    prompt_val = col.get('prompt', '')
                                    max_tokens_val = col.get('max_tokens', 500)
                                    break
                            break
                
                # Show prompt and max_tokens inputs for GenAI Text type
                return [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Text Gen Prompt:", className="form-label fw-bold"),
                            html.Small("Use <Column Name> to reference other column values in your prompt", className="text-muted mb-2 d-block"),
                            dbc.Textarea(
                                id={'type': 'col-prompt', 'op': config_id['op'], 'col': config_id['col']},
                                placeholder="e.g., Write a personalized greeting for <First Name> <Last Name> who works as a <Job Title>",
                                value=prompt_val,
                                rows=3,
                                debounce=True
                            )
                        ], width=12)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            html.Label("Max Tokens:", className="form-label fw-bold"),
                            dcc.Dropdown(
                                id={'type': 'col-max-tokens', 'op': config_id['op'], 'col': config_id['col']},
                                options=[
                                    {'label': '100 tokens', 'value': 100},
                                    {'label': '250 tokens', 'value': 250},
                                    {'label': '500 tokens', 'value': 500},
                                    {'label': '2,500 tokens', 'value': 2500},
                                    {'label': '5,000 tokens', 'value': 5000}
                                ],
                                value=max_tokens_val if max_tokens_val in [100, 250, 500, 2500, 5000] else 500,
                                clearable=False,
                                style={'fontSize': '14px'}
                            )
                        ], width=6)
                    ])
                ]
            elif col_type == 'Country Defined Coordinates':
                # Get existing values for this column
                country_code = 'US'  # Default to US
                coords_only = False  # Default to include all location data
                
                if operations:
                    op_id = config_id['op']
                    col_id = config_id['col']
                    
                    # Find the operation and column to get existing values
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            columns = op['config'].get('columns', [])
                            for col in columns:
                                if col['id'] == col_id:
                                    country_code = col.get('country_code', 'US')
                                    coords_only = col.get('coords_only', False)
                                    break
                            break
                
                # Get country codes list
                try:
                    import pycountry
                    alpha2_codes = [country.alpha_2 for country in pycountry.countries]
                    if 'US' in alpha2_codes:
                        alpha2_codes.remove('US')
                        alpha2_codes.insert(0, 'US')
                    country_options = [{'label': f"{code}", 'value': code} for code in alpha2_codes[:50]]  # Limit to first 50
                except ImportError:
                    # Fallback if pycountry not available
                    country_options = [
                        {'label': 'US', 'value': 'US'},
                        {'label': 'CA', 'value': 'CA'}, 
                        {'label': 'GB', 'value': 'GB'},
                        {'label': 'DE', 'value': 'DE'},
                        {'label': 'FR', 'value': 'FR'},
                        {'label': 'JP', 'value': 'JP'},
                        {'label': 'AU', 'value': 'AU'}
                    ]
                
                # Show country selector and coordinates only option
                return [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Country Code:", className="form-label"),
                            dcc.Dropdown(
                                id={'type': 'col-country-code', 'op': config_id['op'], 'col': config_id['col']},
                                options=country_options,
                                value=country_code,
                                clearable=False,
                                style={'fontSize': '14px'}
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Options:", className="form-label"),
                            dbc.Checklist(
                                id={'type': 'col-coords-only', 'op': config_id['op'], 'col': config_id['col']},
                                options=[{'label': 'Coordinates Only', 'value': 'coords_only'}],
                                value=['coords_only'] if coords_only else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            html.Small([
                                "This will create multiple columns from the base name:", html.Br(),
                                "• {name}_latitude, {name}_longitude", html.Br(),
                                "• If not coordinates only: {name}_city, {name}_country_code, {name}_timezone"
                            ], className="text-muted", style={'fontSize': '12px'})
                        ], width=12)
                    ], className="mt-2")
                ]
            elif col_type == 'Custom Values':
                # Get existing custom values for this column
                custom_values = ['']  # Default: one empty value
                use_weights = False
                custom_weights = [1]  # Default weights
                
                if operations:
                    op_id = config_id['op']
                    col_id = config_id['col']
                    
                    # Find the operation and column to get existing values
                    for op in operations:
                        if op['id'] == op_id and op['type'] == 'tabular':
                            columns = op['config'].get('columns', [])
                            for col in columns:
                                if col['id'] == col_id:
                                    custom_values = col.get('custom_values', [''])
                                    use_weights = col.get('use_weights', False)
                                    custom_weights = col.get('custom_weights', [1] * len(custom_values))
                                    break
                            break
                
                # Ensure we have at least one value
                if not custom_values:
                    custom_values = ['']
                    custom_weights = [1]
                
                # Ensure weights list matches values list length
                while len(custom_weights) < len(custom_values):
                    custom_weights.append(1)
                
                # Build the UI elements
                children = []
                
                # Values section
                children.append(dbc.Row([
                    dbc.Col([
                        html.Label("Custom Values:", className="form-label fw-bold"),
                        html.Div(id={'type': 'custom-values-container', 'op': config_id['op'], 'col': config_id['col']})
                    ], width=12)
                ]))
                
                # Add Value button
                children.append(dbc.Row([
                    dbc.Col([
                        dbc.Button(
                            "+ Add Value",
                            id={'type': 'add-custom-value', 'op': config_id['op'], 'col': config_id['col']},
                            size="sm",
                            color="primary",
                            outline=False,
                            disabled=False
                        )
                    ], width=12)
                ], className="mt-2"))
                
                # Weights checkbox (always render but hide if < 2 values)
                checkbox_style = {} if len(custom_values) >= 2 else {'display': 'none'}
                children.append(dbc.Row([
                    dbc.Col([
                        dbc.Checklist(
                            id={'type': 'use-weights-checkbox', 'op': config_id['op'], 'col': config_id['col']},
                            options=[{'label': ' Use Weights', 'value': 'use_weights'}],
                            value=['use_weights'] if use_weights else [],
                            inline=True,
                            style=checkbox_style
                        )
                    ], width=12)
                ], className="mt-2", style=checkbox_style))
                
                # Add ordered values checkbox for Custom Values
                # Get existing ordered_values state for this column
                ordered_values_state = False
                if operations:
                    for op in operations:
                        if op.get('id') == config_id['op'] and op.get('type') == 'tabular':
                            for col in op.get('config', {}).get('columns', []):
                                if col.get('id') == config_id['col']:
                                    ordered_values_state = col.get('ordered_values', False)
                                    break
                            break
                
                children.append(dbc.Row([
                    dbc.Col([
                        dbc.Checklist(
                            id={'type': 'col-ordered', 'op': config_id['op'], 'col': config_id['col']},
                            options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                            value=['ordered'] if ordered_values_state else [],
                            inline=True,
                            style={'fontSize': '14px'}
                        ),
                        html.Small("When unchecked, values are selected randomly (default). When checked, values are selected in sequence.", 
                                  className="text-muted d-block", style={'fontSize': '12px'})
                    ], width=12)
                ], className="mt-2"))
                
                return children
            else:
                # No additional inputs for First Name or Last Name
                return []
        
        # Custom values container callback - handles dynamic updates to the values input area
        @self.app.callback(
            [Output({'type': 'custom-values-container', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'children'),
             Output({'type': 'use-weights-checkbox', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'style')],
            [Input({'type': 'use-weights-checkbox', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'value'),
             Input('operations-store', 'data')],  # Listen to operations store changes instead
            [State({'type': 'custom-values-container', 'op': dash.dependencies.MATCH, 'col': dash.dependencies.MATCH}, 'id')],
            prevent_initial_call=False
        )
        def update_custom_values_container(weights_checked, operations, container_id):
            if not operations:
                return [], {'display': 'none'}
            
            # Get custom values for this column
            op_id = container_id['op']
            col_id = container_id['col']
            custom_values = ['']
            use_weights = False
            custom_weights = [1]
            
            for op in operations:
                if op['id'] == op_id and op['type'] == 'tabular':
                    columns = op['config'].get('columns', [])
                    for col in columns:
                        if col['id'] == col_id:
                            custom_values = col.get('custom_values', [''])
                            use_weights = bool(weights_checked and 'use_weights' in weights_checked)
                            custom_weights = col.get('custom_weights', [1] * len(custom_values))
                            break
                    break
            
            # Ensure we have at least one value
            if not custom_values:
                custom_values = ['']
                custom_weights = [1]
            
            # Determine checkbox visibility
            checkbox_style = {} if len(custom_values) >= 2 else {'display': 'none'}
            
            return self._create_custom_values_inputs(custom_values, custom_weights, use_weights, op_id, col_id), checkbox_style


        # Add custom value callback
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            Input({'type': 'add-custom-value', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL}, 'n_clicks'),
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def add_custom_value(n_clicks_list, operations):
            print(f"Add custom value callback triggered. n_clicks_list: {n_clicks_list}")
            
            if not operations:
                print("No operations found")
                return dash.no_update
                
            if not any(n_clicks_list):
                print("No clicks detected")
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                print("No context triggered")
                return dash.no_update
            
            print(f"Context triggered: {ctx.triggered}")
            
            try:
                import json
                triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
                print(f"Triggered ID: {triggered_id}")
                triggered_comp = json.loads(triggered_id)
                op_id = triggered_comp['op']
                col_id = triggered_comp['col']
                
                print(f"Looking for op_id: {op_id}, col_id: {col_id}")
                
                # Find the operation and column, add a new custom value
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        columns = op['config'].get('columns', [])
                        for col in columns:
                            if col['id'] == col_id and col.get('data_type') == 'Custom Values':
                                print(f"Found matching column: {col['name']}")
                                custom_values = col.get('custom_values', [''])
                                custom_weights = col.get('custom_weights', [1])
                                
                                print(f"Before: custom_values={custom_values}")
                                
                                # Add new empty value and default weight
                                custom_values.append('')
                                custom_weights.append(1)
                                
                                col['custom_values'] = custom_values
                                col['custom_weights'] = custom_weights
                                
                                print(f"After: custom_values={custom_values}")
                                break
                        break
                
                return operations
            except Exception as e:
                print(f"Error adding custom value: {str(e)}")
                import traceback
                traceback.print_exc()
                return dash.no_update
        
        # Remove custom value callback
        @self.app.callback(
            Output('operations-store', 'data', allow_duplicate=True),
            Input({'type': 'remove-custom-value', 'op': dash.dependencies.ALL, 'col': dash.dependencies.ALL, 'idx': dash.dependencies.ALL}, 'n_clicks'),
            State('operations-store', 'data'),
            prevent_initial_call=True
        )
        def remove_custom_value(n_clicks_list, operations):
            if not operations or not any(n_clicks_list):
                return dash.no_update
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update
            
            try:
                import json
                triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]
                triggered_comp = json.loads(triggered_id)
                op_id = triggered_comp['op']
                col_id = triggered_comp['col']
                idx_to_remove = triggered_comp['idx']
                
                # Find the operation and column, remove the specified value
                for op in operations:
                    if op['id'] == op_id and op['type'] == 'tabular':
                        columns = op['config'].get('columns', [])
                        for col in columns:
                            if col['id'] == col_id and col.get('data_type') == 'Custom Values':
                                custom_values = col.get('custom_values', [''])
                                custom_weights = col.get('custom_weights', [1])
                                
                                # Don't remove if it's the only value
                                if len(custom_values) > 1 and 0 <= idx_to_remove < len(custom_values):
                                    custom_values.pop(idx_to_remove)
                                    if idx_to_remove < len(custom_weights):
                                        custom_weights.pop(idx_to_remove)
                                
                                col['custom_values'] = custom_values
                                col['custom_weights'] = custom_weights
                                break
                        break
                
                return operations
            except Exception as e:
                print(f"Error removing custom value: {str(e)}")
                return dash.no_update
        
        # Download all files callback - only includes files from current batch generation
        @self.app.callback(
            Output('download-files-component', 'data'),
            Input('download-all-files-btn', 'n_clicks'),
            prevent_initial_call=True
        )
        def download_all_files(n_clicks):
            if not n_clicks:
                return dash.no_update
            
            try:
                print(f"🔽 Download button clicked, starting zip creation...")
                
                # Create timestamp for zip filename
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                
                # Create zip file with files from current batch generation only
                # (completed_items is reset at start of each batch)
                # Smart approach: download missing files to local first, then zip locally
                zip_info = self._create_zip_from_local_files(timestamp)
                
                if zip_info and zip_info['local_path']:
                    print(f"📦 Zip created successfully: {zip_info['filename']} ({zip_info['file_count']} files)")
                    
                    # Return file for download
                    return dcc.send_file(zip_info['local_path'], filename=zip_info['filename'])
                else:
                    print("❌ Failed to create zip file")
                    return dash.no_update
                    
            except Exception as e:
                print(f"❌ ERROR in download callback: {str(e)}")
                import traceback
                print(f"Full traceback: {traceback.format_exc()}")
                return dash.no_update

        # NOTE: Row count validation is now handled directly in the main operations callback above
        # to avoid race conditions between multiple callbacks

    def _create_custom_values_inputs(self, custom_values, custom_weights, use_weights, op_id, col_id):
        """Create input fields for custom values and optional weights."""
        children = []
        
        for i, (value, weight) in enumerate(zip(custom_values, custom_weights)):
            row_children = []
            
            # Value input
            if use_weights:
                # Value input takes 8 columns, weight takes 3, remove button takes 1
                row_children.append(dbc.Col([
                    dbc.Input(
                        id={'type': 'custom-value-input', 'op': op_id, 'col': col_id, 'idx': i},
                        placeholder=f"Value {i+1}",
                        value=value,
                        debounce=True
                    )
                ], width=8))
                
                # Weight input
                row_children.append(dbc.Col([
                    dbc.Input(
                        id={'type': 'custom-weight-input', 'op': op_id, 'col': col_id, 'idx': i},
                        type="number",
                        placeholder="Weight",
                        value=weight,
                        min=0.1,
                        step=0.1,
                        debounce=True
                    )
                ], width=3))
            else:
                # Value input takes 11 columns, remove button takes 1
                row_children.append(dbc.Col([
                    dbc.Input(
                        id={'type': 'custom-value-input', 'op': op_id, 'col': col_id, 'idx': i},
                        placeholder=f"Value {i+1}",
                        value=value,
                        debounce=True
                    )
                ], width=11))
            
            # Remove button (only show if more than 1 value)
            if len(custom_values) > 1:
                row_children.append(dbc.Col([
                    dbc.Button(
                        "×",
                        id={'type': 'remove-custom-value', 'op': op_id, 'col': col_id, 'idx': i},
                        size="sm",
                        color="danger",
                        outline=True
                    )
                ], width=1))
            else:
                # Empty column to maintain alignment
                row_children.append(dbc.Col([], width=1))
            
            children.append(dbc.Row(row_children, className="mb-2"))
        
        return children

    def _create_operation_card(self, operation):
        """Create a card for configuring an operation."""
        op_id = operation['id']
        op_type = operation['type']
        configured = operation.get('configured', False)
        config = operation.get('config', {})
        
        # Create operation-specific configuration inputs
        if op_type == 'pdf':
            config_inputs = [
                html.Label("Document Name:", className="form-label fw-bold"),
                dbc.Input(
                    id={'type': 'pdf-doc-name', 'index': op_id},
                    placeholder="Enter document name (e.g., Employee Handbook, Privacy Policy)",
                    value=config.get('doc_name', ''),
                    debounce=True,
                    className="mb-3"
                ),
                html.Label("Document Type:", className="form-label fw-bold"),
                dcc.Dropdown(
                    id={'type': 'pdf-doc-type', 'index': op_id},
                    options=[
                        {'label': 'Policy Guide', 'value': 'policy_guide'},
                        {'label': 'Customer Correspondence', 'value': 'customer_correspondence'},
                        {'label': 'Customer Profile', 'value': 'customer_profile'},
                        {'label': 'Custom Document', 'value': 'custom_document'}
                    ],
                    value=config.get('doc_type', 'policy_guide'),
                    className="mb-3"
                ),
                html.Label("Description:", className="form-label fw-bold"),
                dbc.Textarea(
                    id={'type': 'pdf-description', 'index': op_id},
                    placeholder="Describe the PDF document you want to generate...",
                    value=config.get('description', ''),
                    rows=3,
                    debounce=True,
                    className="mb-3"
                ),
                html.Label("Number of Documents:", className="form-label fw-bold"),
                dcc.Slider(
                    id={'type': 'pdf-count', 'index': op_id},
                    min=1,
                    max=10,
                    step=1,
                    value=config.get('count', 1),
                    marks={i: str(i) for i in range(1, 11)},
                    className="mb-3"
                ),
                html.Div([
                    dbc.Checklist(
                        id={'type': 'pdf-include-images', 'index': op_id},
                        options=[
                            {'label': ' Include AI-Generated Images', 'value': 'include_images'}
                        ],
                        value=['include_images'] if config.get('include_images', False) else [],
                        inline=True,
                        className="mb-2"
                    ),
                    html.Small("Generate 3 contextual images based on document content", 
                              className="text-muted d-block")
                ], className="mb-3")
            ]
        elif op_type == 'text':
            config_inputs = [
                html.Label("Document Name:", className="form-label fw-bold"),
                dbc.Input(
                    id={'type': 'text-doc-name', 'index': op_id},
                    placeholder="Enter document name (e.g., Employee Handbook, Privacy Policy)",
                    value=config.get('doc_name', ''),
                    debounce=True,
                    className="mb-3"
                ),
                html.Label("Document Type:", className="form-label fw-bold"),
                dcc.Dropdown(
                    id={'type': 'text-doc-type', 'index': op_id},
                    options=[
                        {'label': 'Policy Guide', 'value': 'policy_guide'},
                        {'label': 'Customer Correspondence', 'value': 'customer_correspondence'},
                        {'label': 'Customer Profile', 'value': 'customer_profile'},
                        {'label': 'Custom Document', 'value': 'custom_document'}
                    ],
                    value=config.get('doc_type', 'policy_guide'),
                    className="mb-3"
                ),
                html.Label("File Format:", className="form-label fw-bold"),
                dcc.Dropdown(
                    id={'type': 'text-format', 'index': op_id},
                    options=[
                        {'label': 'Text File (.txt)', 'value': 'txt'},
                        {'label': 'Word Document (.docx)', 'value': 'docx'}
                    ],
                    value=config.get('file_format', 'txt'),
                    className="mb-3"
                ),
                html.Label("Description:", className="form-label fw-bold"),
                dbc.Textarea(
                    id={'type': 'text-description', 'index': op_id},
                    placeholder="Describe the text document you want to generate...",
                    value=config.get('description', ''),
                    rows=3,
                    debounce=True,
                    className="mb-3"
                ),
                html.Label("Number of Documents:", className="form-label fw-bold"),
                dcc.Slider(
                    id={'type': 'text-count', 'index': op_id},
                    min=1,
                    max=10,
                    step=1,
                    value=config.get('count', 1),
                    marks={i: str(i) for i in range(1, 11)},
                    className="mb-3"
                )
            ]
        elif op_type == 'tabular':
            config_inputs = [
                html.Label("Table Name:", className="form-label fw-bold"),
                dbc.Input(
                    id={'type': 'tabular-name', 'index': op_id},
                    placeholder="Enter table name...",
                    value=config.get('table_name', ''),
                    debounce=True,
                    className="mb-3"
                ),
                html.Label("Number of Rows:", className="form-label fw-bold"),
                dbc.Input(
                    id={'type': 'tabular-rows', 'index': op_id},
                    type="number",
                    min=1,
                    # NO MAX ATTRIBUTE - let our Python code handle validation
                    step=1,
                    value=config.get('row_count', 1000),
                    placeholder="Enter number of rows (any size - validation applied in backend)",
                    debounce=True,  # Add debounce back for smooth typing experience
                    className="mb-1",
                    # Add input validation attributes
                    pattern="[0-9]*",  # Only allow numbers
                    inputMode="numeric"  # Show numeric keypad on mobile
                ),
                html.Small(
                    "Maximum 1,000,000 rows normally, 1,000 rows if GenAI Text columns are present",
                    className="text-muted mb-3 d-block",
                    style={'fontSize': '12px'}
                ),
                html.Div([
                    html.Div([
                        html.Label("Columns:", className="form-label fw-bold"),
                        dbc.Button(
                            "Add Column",
                            id={'type': 'add-column', 'index': op_id},
                            color="primary",
                            size="sm",
                            className="float-end"
                        )
                    ], className="d-flex justify-content-between align-items-center mb-3"),
                    html.Div(
                        id={'type': 'columns-container', 'index': op_id},
                        children=self._create_column_cards(config.get('columns', []), op_id),
                        className="mb-3"
                    )
                ])
            ]
        else:
            config_inputs = [html.P("Unknown operation type")]
        
        # Status indicator
        if configured:
            status_color = "success"
            status_icon = "✅"
            status_text = "Configured"
        else:
            status_color = "warning"
            status_icon = "⚠️"
            status_text = "Needs Configuration"
        
        return dbc.Card([
            dbc.CardHeader([
                html.Div([
                    html.H5(f"{self._format_data_type(op_type)}", className="mb-0"),
                    html.Div([
                        dbc.Badge(f"{status_icon} {status_text}", color=status_color, className="me-2"),
                        dbc.Button("Remove", id={'type': 'remove-op', 'index': op_id}, color="danger", size="sm")
                    ])
                ], className="d-flex justify-content-between align-items-center")
            ]),
            dbc.CardBody(config_inputs)
        ], className="mb-3")

    def _create_column_cards(self, columns, op_id):
        """Create cards for column configuration within a tabular operation."""
        if not columns:
            return []
        
        column_cards = []
        for i, col in enumerate(columns):
            col_id = col.get('id', f"col_{i}")
            col_name = col.get('name', '')
            col_type = col.get('data_type', 'Integer')
            
            # Type-specific inputs based on column type
            if col_type == 'Integer':
                type_inputs = [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Min Value:", className="form-label"),
                            dbc.Input(
                                id={'type': 'col-min', 'op': op_id, 'col': col_id},
                                type="number",
                                value=col.get('min_value', 1),
                                size="sm",
                                debounce=True
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Max Value:", className="form-label"),
                            dbc.Input(
                                id={'type': 'col-max', 'op': op_id, 'col': col_id},
                                type="number",
                                value=col.get('max_value', 100),
                                size="sm",
                                debounce=True
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dbc.Checklist(
                                id={'type': 'col-ordered', 'op': op_id, 'col': col_id},
                                options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                                value=['ordered'] if col.get('ordered_values', False) else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            ),
                            html.Small("When unchecked, values are randomized (default). When checked, values are generated in sequence.", 
                                      className="text-muted d-block", style={'fontSize': '12px'})
                        ], width=12)
                    ])
                ]
            elif col_type == 'Date':
                type_inputs = [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Min Date:", className="form-label"),
                            dcc.DatePickerSingle(
                                id={'type': 'col-min-date', 'op': op_id, 'col': col_id},
                                date=col.get('min_date', '2020-01-01'),
                                display_format='YYYY-MM-DD',
                                style={'width': '100%', 'fontSize': '14px'}
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Max Date:", className="form-label"),
                            dcc.DatePickerSingle(
                                id={'type': 'col-max-date', 'op': op_id, 'col': col_id},
                                date=col.get('max_date', '2024-12-31'),
                                display_format='YYYY-MM-DD',
                                style={'width': '100%', 'fontSize': '14px'}
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dbc.Checklist(
                                id={'type': 'col-ordered', 'op': op_id, 'col': col_id},
                                options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                                value=['ordered'] if col.get('ordered_values', False) else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            ),
                            html.Small("When unchecked, dates are randomized (default). When checked, dates are generated in sequence.", 
                                      className="text-muted d-block", style={'fontSize': '12px'})
                        ], width=12)
                    ])
                ]
            elif col_type == 'Country Defined Coordinates':
                # Get country codes list
                try:
                    import pycountry
                    alpha2_codes = [country.alpha_2 for country in pycountry.countries]
                    if 'US' in alpha2_codes:
                        alpha2_codes.remove('US')
                        alpha2_codes.insert(0, 'US')
                    country_options = [{'label': f"{code}", 'value': code} for code in alpha2_codes[:50]]  # Limit to first 50
                except ImportError:
                    # Fallback if pycountry not available
                    country_options = [
                        {'label': 'US', 'value': 'US'},
                        {'label': 'CA', 'value': 'CA'}, 
                        {'label': 'GB', 'value': 'GB'},
                        {'label': 'DE', 'value': 'DE'},
                        {'label': 'FR', 'value': 'FR'},
                        {'label': 'JP', 'value': 'JP'},
                        {'label': 'AU', 'value': 'AU'}
                    ]
                
                type_inputs = [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Country Code:", className="form-label"),
                            dcc.Dropdown(
                                id={'type': 'col-country-code', 'op': op_id, 'col': col_id},
                                options=country_options,
                                value=col.get('country_code', 'US'),
                                clearable=False,
                                style={'fontSize': '14px'}
                            )
                        ], width=6),
                        dbc.Col([
                            html.Label("Options:", className="form-label"),
                            dbc.Checklist(
                                id={'type': 'col-coords-only', 'op': op_id, 'col': col_id},
                                options=[{'label': 'Coordinates Only', 'value': 'coords_only'}],
                                value=['coords_only'] if col.get('coords_only', False) else [],
                                inline=True,
                                style={'fontSize': '14px'}
                            )
                        ], width=6)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            html.Small([
                                "This will create multiple columns from the base name:", html.Br(),
                                "• {name}_latitude, {name}_longitude", html.Br(),
                                "• If not coordinates only: {name}_city, {name}_country_code, {name}_timezone"
                            ], className="text-muted", style={'fontSize': '12px'})
                        ], width=12)
                    ], className="mt-2")
                ]
            elif col_type == 'GenAI Text':
                type_inputs = [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Text Gen Prompt:", className="form-label fw-bold"),
                            html.Small("Use <Column Name> to reference other column values in your prompt", className="text-muted mb-2 d-block"),
                            dbc.Textarea(
                                id={'type': 'col-prompt', 'op': op_id, 'col': col_id},
                                placeholder="e.g., Write a personalized greeting for <First Name> <Last Name> who works as a <Job Title>",
                                value=col.get('prompt', ''),
                                rows=3,
                                debounce=True
                            )
                        ], width=12)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            html.Label("Max Tokens:", className="form-label fw-bold"),
                            dcc.Dropdown(
                                id={'type': 'col-max-tokens', 'op': op_id, 'col': col_id},
                                options=[
                                    {'label': '100 tokens', 'value': 100},
                                    {'label': '250 tokens', 'value': 250},
                                    {'label': '500 tokens', 'value': 500},
                                    {'label': '2,500 tokens', 'value': 2500},
                                    {'label': '5,000 tokens', 'value': 5000}
                                ],
                                value=col.get('max_tokens', 500) if col.get('max_tokens', 500) in [100, 250, 500, 2500, 5000] else 500,
                                clearable=False,
                                style={'fontSize': '14px'}
                            )
                        ], width=6)
                    ])
                ]
            elif col_type == 'Custom Values':
                # For Custom Values, create the container that will be populated by callback
                custom_values = col.get('custom_values', [''])
                use_weights = col.get('use_weights', False)
                custom_weights = col.get('custom_weights', [1])
                
                type_inputs = [
                    dbc.Row([
                        dbc.Col([
                            html.Label("Custom Values:", className="form-label fw-bold"),
                            html.Div(
                                self._create_custom_values_inputs(custom_values, custom_weights, use_weights, op_id, col_id),
                                id={'type': 'custom-values-container', 'op': op_id, 'col': col_id}
                            )
                        ], width=12)
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dbc.Button(
                                "+ Add Value",
                                id={'type': 'add-custom-value', 'op': op_id, 'col': col_id},
                                size="sm",
                                color="primary",
                                outline=False,
                                disabled=False
                            )
                        ], width=12)
                    ], className="mt-2"),
                ]
                
                # Add weights checkbox (always render but hide if < 2 values)
                checkbox_style = {} if len(custom_values) >= 2 else {'display': 'none'}
                type_inputs.append(dbc.Row([
                    dbc.Col([
                        dbc.Checklist(
                            id={'type': 'use-weights-checkbox', 'op': op_id, 'col': col_id},
                            options=[{'label': ' Use Weights', 'value': 'use_weights'}],
                            value=['use_weights'] if use_weights else [],
                            inline=True,
                            style=checkbox_style
                        )
                    ], width=12)
                ], className="mt-2", style=checkbox_style))
                
                # Add ordered values checkbox
                type_inputs.append(dbc.Row([
                    dbc.Col([
                        dbc.Checklist(
                            id={'type': 'col-ordered', 'op': op_id, 'col': col_id},
                            options=[{'label': 'Ordered Values', 'value': 'ordered'}],
                            value=['ordered'] if col.get('ordered_values', False) else [],
                            inline=True,
                            style={'fontSize': '14px'}
                        ),
                        html.Small("When unchecked, values are randomized (default). When checked, values are selected in sequence.", 
                                  className="text-muted d-block", style={'fontSize': '12px'})
                    ], width=12)
                ], className="mt-2"))
            else:
                # No additional inputs for First Name or Last Name
                type_inputs = []
            
            card = dbc.Card([
                dbc.CardBody([
                    dbc.Row([
                        dbc.Col([
                            html.Label("Column Name:", className="form-label"),
                            dbc.Input(
                                id={'type': 'col-name', 'op': op_id, 'col': col_id},
                                value=col_name,
                                size="sm",
                                debounce=True,
                                placeholder="Enter column name..."
                            )
                        ], width=4),
                        dbc.Col([
                            html.Label("Data Type:", className="form-label"),
                            dcc.Dropdown(
                                id={'type': 'col-type', 'op': op_id, 'col': col_id},
                                options=[
                                    {'label': 'Integer', 'value': 'Integer'},
                                    {'label': 'Date', 'value': 'Date'},
                                    {'label': 'First Name', 'value': 'First Name'},
                                    {'label': 'Last Name', 'value': 'Last Name'},
                                    {'label': 'Email', 'value': 'Email'},
                                    {'label': 'Phone Number', 'value': 'Phone Number'},
                                    {'label': 'Address', 'value': 'Address'},
                                    {'label': 'City', 'value': 'City'},
                                    {'label': 'Country', 'value': 'Country'},
                                    {'label': 'Country Code', 'value': 'Country Code'},
                                    {'label': 'Postcode', 'value': 'Postcode'},
                                    {'label': 'Street Address', 'value': 'Street Address'},
                                    {'label': 'Company', 'value': 'Company'},
                                    {'label': 'Credit Card Number', 'value': 'Credit Card Number'},
                                    {'label': 'Credit Card Provider', 'value': 'Credit Card Provider'},
                                    {'label': 'Latitude', 'value': 'Latitude'},
                                    {'label': 'Longitude', 'value': 'Longitude'},
                                    {'label': 'Country Defined Coordinates', 'value': 'Country Defined Coordinates'},
                                    {'label': 'GenAI Text', 'value': 'GenAI Text'},
                                    {'label': 'Custom Values', 'value': 'Custom Values'}
                                ],
                                value=col_type,
                                style={'font-size': '14px'}
                            )
                        ], width=6),
                        dbc.Col([
                            dbc.Button(
                                "Remove",
                                id={'type': 'remove-column', 'op': op_id, 'col': col_id},
                                color="danger",
                                size="sm",
                                className="mt-4"
                            )
                        ], width=2)
                    ], className="mb-2"),
                    html.Div(type_inputs, id={'type': 'col-config', 'op': op_id, 'col': col_id})
                ])
            ], className="mb-2")
            
            column_cards.append(card)
        
        return column_cards

    def _generate_batch_background(self, operations, company_name, company_sector):
        """Generate multiple operations in background thread."""
        try:
            for i, operation in enumerate(operations):
                # Update progress
                self.generation_state['current_operation'] = i + 1
                op_type = operation['type']
                self.generation_state['current_step'] = f"Processing {self._format_data_type(op_type)} operation..."
                
                # Generate based on operation type and configuration
                if op_type == 'pdf':
                    items = self._process_pdf_operation(operation, company_name, company_sector)
                elif op_type == 'text':
                    items = self._process_text_operation(operation, company_name, company_sector)
                elif op_type == 'tabular':
                    items = self._process_tabular_operation(operation, company_name, company_sector)
                else:
                    print(f"Unknown operation type: {op_type}")
                    continue
                
                # Add items to completed list
                if isinstance(items, list):
                    self.generation_state['completed_items'].extend(items)
                else:
                    self.generation_state['completed_items'].append(items)
                
                # Small delay between operations
                if i < len(operations) - 1:
                    self.generation_state['current_step'] = f"Preparing next operation..."
                    time.sleep(0.5)
            
            # Mark as complete
            self.generation_state['active'] = False
            self.generation_state['current_step'] = 'Batch complete!'
            
        except Exception as e:
            print(f"Error in batch generation: {str(e)}")
            self.generation_state['active'] = False
            self.generation_state['error'] = str(e)

    def _process_pdf_operation(self, operation, company_name, company_sector):
        """Process a PDF generation operation."""
        config = operation.get('config', {})
        doc_type = config.get('doc_type', 'policy_guide')
        doc_name = config.get('doc_name', '')
        description = config.get('description', 'Sample document')
        count = config.get('count', 1)
        include_images = config.get('include_images', False)
        
        items = []
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        for i in range(count):
            enhanced_description = f"For {company_name} (a {company_sector} company): {description}"
            item = self._generate_pdf_item(enhanced_description, company_name, company_sector, f"{timestamp}_{i+1}", doc_type, include_images, doc_name)
            items.append(item)
        
        return items

    def _process_text_operation(self, operation, company_name, company_sector):
        """Process a text generation operation."""
        config = operation.get('config', {})
        doc_type = config.get('doc_type', 'policy_guide')
        doc_name = config.get('doc_name', '')
        description = config.get('description', 'Sample document')
        file_format = config.get('file_format', 'txt')
        count = config.get('count', 1)
        
        items = []
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        for i in range(count):
            enhanced_description = f"For {company_name} (a {company_sector} company): {description}"
            item = self._generate_text_item(enhanced_description, company_name, company_sector, f"{timestamp}_{i+1}", doc_type, file_format, doc_name)
            items.append(item)
        
        return items

    def _process_tabular_operation(self, operation, company_name, company_sector):
        """Process a tabular data generation operation using dbldatagen."""
        config = operation.get('config', {})
        table_name = config.get('table_name', 'sample_table')
        row_count = config.get('row_count', 1000)
        columns = config.get('columns', [])
        
        if not columns:
            # If no columns configured, skip this operation
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        return self._generate_tabular_item(table_name, row_count, columns, company_name, company_sector, timestamp)

    def _generate_documents_background(self, doc_type, description, count, doc_name=''):
        """Generate documents in background thread with progress updates."""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            for i in range(count):
                # Update progress
                self.generation_state['current_doc'] = i + 1
                self.generation_state['current_step'] = f"Generating content for document {i + 1}..."
                
                # Generate content using the serving endpoint
                content = self._generate_document_content(doc_type, description, i + 1)
                
                # Update progress
                self.generation_state['current_step'] = f"Creating PDF for document {i + 1}..."
                
                # Create filename using document name if provided, otherwise fallback to doc_type
                if doc_name:
                    # Sanitize document name for filename
                    safe_name = "".join(c for c in doc_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    safe_name = safe_name.replace(' ', '_')
                    filename = f"{safe_name}_{timestamp}_{i+1:02d}.pdf"
                else:
                    filename = f"{doc_type}_{timestamp}_{i+1:02d}.pdf"
                
                # Generate PDF
                pdf_path = self._create_pdf(content, filename, doc_type)
                
                # Update progress and save to volume
                self.generation_state['current_step'] = f"Saving document {i + 1} to volume..."
                
                # Save to volume and check success
                save_success = self._save_to_volume(pdf_path, filename)
                
                file_info = {
                    'filename': filename,
                    'path': pdf_path,
                    'doc_type': doc_type,
                    'size': os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0,
                    'saved_to_volume': save_success
                }
                
                self.generation_state['completed_files'].append(file_info)
                
                # Small delay between generations
                if i < count - 1:  # Don't delay after the last document
                    self.generation_state['current_step'] = f"Preparing next document..."
                    time.sleep(0.5)
            
            # Mark as complete
            self.generation_state['active'] = False
            self.generation_state['current_step'] = 'Complete!'
            
        except Exception as e:
            print(f"Error in background generation: {str(e)}")
            self.generation_state['active'] = False
            self.generation_state['error'] = str(e)

    def _generate_document_content(self, doc_type, description, doc_number):
        """Generate document content using the serving endpoint."""
        
        # Create document-specific prompts
        doc_prompts = {
            'policy_guide': f"""Create a comprehensive policy guide document with the following characteristics:
{description}

This should be document #{doc_number} in a series. Make it realistic and detailed with:
- Clear policy sections and subsections
- Specific procedures and guidelines
- Professional formatting with headers
- Realistic company policies and procedures
- Compliance and regulatory information where appropriate

Generate a complete document with multiple sections, not just an outline.""",
            
            'customer_correspondence': f"""Create a realistic customer correspondence document with these characteristics:
{description}

This should be document #{doc_number} in a series. Include:
- Realistic customer information (use fictional names and details)
- Professional business communication tone
- Specific details about products, services, or issues
- Appropriate formatting for business correspondence
- Dates, reference numbers, and other realistic details

Generate a complete correspondence document, not just a template.""",
            
            'customer_profile': f"""Create a detailed customer profile document with these characteristics:
{description}

This should be document #{doc_number} in a series. Include:
- Comprehensive customer demographics and information
- Purchase history and preferences
- Account details and status
- Contact information and communication preferences
- Notes and interactions history
- Risk assessment or credit information where relevant

Generate a complete customer profile with realistic data (use fictional information).""",

            'custom_document': f"""Create a custom document with these characteristics:
{description}

This should be document #{doc_number} in a series. Make it realistic and well-structured based on the requirements provided in the description. Include:
- Appropriate formatting and structure for the document type
- Professional presentation and language
- Realistic details and content
- Clear organization and flow
- Complete content, not just an outline or template

Generate a complete document that fulfills the specific requirements described."""
        }
        
        prompt = doc_prompts.get(doc_type, f"Create a {doc_type} document with these characteristics: {description}")
        
        messages = [
            {"role": "system", "content": "You are a professional document generator. Create realistic, detailed documents based on the user's specifications. Use fictional but realistic data. Make the documents comprehensive and well-structured. IMPORTANT: Return ONLY the document content itself, no reasoning, no metadata, no explanations - just the complete document text ready for PDF generation."},
            {"role": "user", "content": prompt + "\n\nIMPORTANT: Respond with ONLY the complete document content. Do not include any reasoning, metadata, or explanations. Start directly with the document title and content."}
        ]
        
        try:
            response = query_endpoint(self.endpoint_name, messages, max_tokens=2048)
            
            # Debug logging to understand response structure
            print(f"DEBUG: Raw response type: {type(response)}")
            if isinstance(response, dict):
                print(f"DEBUG: Response keys: {list(response.keys())}")
                if 'type' in response:
                    print(f"DEBUG: Response type field: {response['type']}")
            
            # Extract content using robust logic to handle different response types
            content = self._extract_content_safely(response)
            
            # Ensure content is always a string
            if isinstance(content, list):
                # If content is a list, join it into a string
                content = '\n\n'.join(str(item) for item in content)
            elif not isinstance(content, str):
                content = str(content)
            
            return content
                
        except Exception as e:
            print(f"Error generating content: {str(e)}")
            # Fallback content
            return f"Sample {self._format_doc_type(doc_type)} Document #{doc_number}\n\n{description}\n\nThis is a placeholder document generated due to an API error."

    def _create_pdf(self, content, filename, doc_type):
        """Create a PDF file from the generated content."""
        
        # Ensure the volume directory exists (create locally for now)
        local_dir = "./generated_documents"
        os.makedirs(local_dir, exist_ok=True)
        
        pdf_path = os.path.join(local_dir, filename)
        
        # Create PDF with enhanced formatting (matching _create_pdf_with_images)
        doc = SimpleDocTemplate(
            pdf_path, 
            pagesize=letter,
            leftMargin=1*inch,
            rightMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        styles = getSampleStyleSheet()
        
        # Enhanced custom styles for professional appearance (same as with images)
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=36,
            spaceBefore=12,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold',
            textColor='#2E3440'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=18,
            spaceBefore=24,
            alignment=0,  # Left alignment
            fontName='Helvetica-Bold',
            textColor='#3B4252'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=14,
            spaceBefore=2,
            alignment=4,  # Justified alignment
            fontName='Helvetica',
            leftIndent=0,
            rightIndent=0,
            firstLineIndent=0,
            leading=16,  # Line spacing
            textColor='#2E3440'
        )
        
        # Special style for bullet points and lists
        list_style = ParagraphStyle(
            'CustomList',
            parent=body_style,
            fontSize=12,
            spaceAfter=8,
            spaceBefore=4,
            leftIndent=20,
            bulletIndent=10,
            leading=15
        )
        
        # Build document
        story = []
        
        # Add document header with company info and date (same as with images)
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")
        
        header_info = f"Generated on {current_date}"
        header_style = ParagraphStyle(
            'HeaderInfo',
            parent=styles['Normal'],
            fontSize=9,
            alignment=2,  # Right alignment
            textColor='#5E81AC',
            spaceAfter=20
        )
        story.append(Paragraph(header_info, header_style))
        
        # Title with better spacing
        title = f"{self._format_doc_type(doc_type)} Document"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Sanitize content for PDF generation with enhanced formatting
        content = self._sanitize_content_for_pdf(content)
        
        # Content - split by paragraphs and create PDF elements with enhanced detection
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                try:
                    # Clean the paragraph text for ReportLab
                    para_text = para.strip()
                    
                    # Escape special characters that might cause ReportLab issues
                    para_text = para_text.replace('&', '&amp;')
                    para_text = para_text.replace('<', '&lt;')
                    para_text = para_text.replace('>', '&gt;')
                    
                    # Enhanced paragraph handling (same logic as with images)
                    import re
                    
                    # Table of Contents
                    if 'TABLE OF CONTENTS' in para_text.upper():
                        toc_style = ParagraphStyle(
                            'TOC_Header',
                            parent=subtitle_style,
                            fontSize=16,
                            alignment=1,  # Center
                            spaceAfter=20,
                            spaceBefore=20
                        )
                        story.append(Paragraph(para_text, toc_style))
                        story.append(Spacer(1, 0.2*inch))
                    
                    # Numbered sections
                    elif re.match(r'^\d+\.\s+[A-Z]', para_text):
                        story.append(Paragraph(para_text, subtitle_style))
                        story.append(Spacer(1, 0.15*inch))
                    
                    # Document metadata
                    elif any(keyword in para_text for keyword in ['Effective Date:', 'Document #', 'Page ', 'Revision #:']):
                        metadata_style = ParagraphStyle(
                            'Metadata',
                            parent=body_style,
                            fontSize=10,
                            textColor='#5E81AC',
                            alignment=1,  # Center
                            spaceBefore=8,
                            spaceAfter=16
                        )
                        story.append(Paragraph(para_text, metadata_style))
                        story.append(Spacer(1, 0.12*inch))
                    
                    # Bullet points and lists
                    elif para_text.startswith('•') or para_text.startswith('-') or para_text.startswith('*') or para_text.strip().startswith('•'):
                        story.append(Paragraph(para_text, list_style))
                        story.append(Spacer(1, 0.08*inch))
                    
                    # Section headers
                    elif len(para_text) < 80 and (para_text.endswith(':') or para_text.isupper()):
                        story.append(Paragraph(para_text, subtitle_style))
                        story.append(Spacer(1, 0.15*inch))
                    
                    # Regular body text
                    else:
                        story.append(Paragraph(para_text, body_style))
                        story.append(Spacer(1, 0.08*inch))
                        
                except Exception as e:
                    # If paragraph creation fails, add as plain text
                    print(f"Warning: Could not create paragraph, adding as plain text: {str(e)}")
                    # Convert to safe text and add as simple paragraph
                    safe_text = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    try:
                        story.append(Paragraph(safe_text, body_style))
                        story.append(Spacer(1, 0.08*inch))
                    except:
                        # Last resort: skip this paragraph
                        print(f"Skipping problematic paragraph: {para[:100]}...")
                        continue
        
        # Build PDF
        doc.build(story)
        
        return pdf_path
    
    def _generate_image_prompts(self, pdf_content):
        """Generate 3 contextual image prompts based on PDF content using robust parsing."""
        try:
            from mlflow.deployments import get_deploy_client
            
            # Create a prompt to generate image prompts with clear formatting
            system_prompt = """Based on the provided document content, generate exactly 3 descriptive image prompts that would enhance the document when included as illustrations. Each prompt should:

1. Be clear and specific for image generation
2. Relate to different sections or concepts in the document
3. Be appropriate for a professional document
4. Be realistic and achievable by an image generator

Format your response in one of these ways:
- As a JSON array: ["prompt1", "prompt2", "prompt3"]
- As numbered lines:
1. prompt1
2. prompt2  
3. prompt3

Provide exactly 3 prompts, each on its own line or in a clear format."""

            user_prompt = f"Document content:\n\n{pdf_content}\n\nGenerate 3 image prompts for this document."

            # Prepare messages for the endpoint
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]

            # Query the LLM for image prompts - no responseFormat needed
            print("🧠 Requesting image prompts from databricks-meta-llama-3-3-70b-instruct...")
            response = get_deploy_client('databricks').predict(
                endpoint="databricks-meta-llama-3-3-70b-instruct",
                inputs={
                    "messages": messages, 
                    "max_tokens": 800
                }
            )
            
            print(f"🔍 Response received: {type(response)}")
            
            if response:
                # Extract content from response
                try:
                    # Handle both direct response and wrapped response formats
                    if isinstance(response, dict):
                        if "choices" in response and len(response["choices"]) > 0:
                            # OpenAI-style response format
                            content = response["choices"][0]["message"]["content"]
                        elif "messages" in response and len(response["messages"]) > 0:
                            # Databricks agent format
                            content = response["messages"][-1]["content"]
                        else:
                            # Direct content
                            content = str(response)
                    else:
                        content = str(response)
                    
                    print(f"📝 Extracted content for parsing...")
                    
                    # Use robust parsing to extract image prompts
                    parsed_prompts = self._extract_image_prompts_robustly(content)
                    
                    if parsed_prompts and len(parsed_prompts) == 3:
                        print(f"🎯 Successfully extracted 3 image prompts using robust parsing")
                        return parsed_prompts
                    else:
                        print(f"⚠️ Robust parsing returned {len(parsed_prompts) if parsed_prompts else 0} prompts")
                        
                except Exception as e:
                    print(f"⚠️ Error processing LLM response: {e}")
                    print(f"Response content: {str(response)[:200]}...")
            else:
                print("⚠️ No response from LLM for image prompt generation")
                
        except Exception as e:
            print(f"❌ Error in image prompt generation: {str(e)}")
            
        # Fallback prompts in case of any error
        print("🔄 All parsing strategies failed, using fallback image prompts")
        return [
            "A professional business document illustration with clean modern design",
            "Office workers collaborating on important company policies and procedures", 
            "A corporate meeting room with presentations and documentation on display"
        ]
    
    def _generate_images(self, image_prompts):
        """Generate images from prompts using the Databricks serving endpoint."""
        generated_images = []
        
        try:
            from openai import OpenAI
            from databricks.sdk import WorkspaceClient
            import base64
            
            # Get API key from Databricks secrets
            print("🔐 Retrieving API key from Databricks secrets...")
            
            def get_secret(scope, key):
                try:
                    w = WorkspaceClient()
                    secret_response = w.secrets.get_secret(scope=scope, key=key)
                    decoded_secret = base64.b64decode(secret_response.value).decode('utf-8')
                    return decoded_secret
                except Exception as e:
                    print(f"❌ Secret not found or inaccessible: {e}")
                    return None

            scope_name = "css_secrets"
            secret_key = "image_gen_key"
            api_key = get_secret(scope_name, secret_key)
            
            if api_key:
                print("✅ API key retrieved from Databricks secrets")
            else:
                print("❌ Failed to retrieve API key from secrets, trying environment fallback")
                api_key = os.environ.get('DATABRICKS_TOKEN')
                if api_key:
                    print("✅ Using API key from environment variable")
                else:
                    print("❌ No API key found in secrets or DATABRICKS_TOKEN environment variable")
                    return []

            client = OpenAI(
                api_key=api_key,
                base_url="https://e2-demo-field-eng.cloud.databricks.com/serving-endpoints"
            )

            
            for i, prompt in enumerate(image_prompts):
                try:
                    print(f"🎨 Generating image {i+1}/3: {prompt[:50]}...")

                    response = client.images.generate(
                        model="databricks-shutterstock-imageai",
                        prompt=prompt  # Uses the actual contextual prompt, not hardcoded test
                    )
                    
                    # Debug: Log the response structure
                    print(f"🔍 Response type: {type(response)}")
                    print(f"🔍 Response attributes: {dir(response)}")
                    if hasattr(response, 'data'):
                        print(f"🔍 Response.data type: {type(response.data)}")
                        print(f"🔍 Response.data length: {len(response.data)}")
                        if len(response.data) > 0:
                            print(f"🔍 First data item type: {type(response.data[0])}")
                            print(f"🔍 First data item attributes: {dir(response.data[0])}")
                    
                    if response and hasattr(response, 'data') and len(response.data) > 0:
                        # Try different ways to extract image data
                        first_item = response.data[0]
                        
                        # Try multiple possible response structures
                        image_data = None
                        if hasattr(first_item, 'b64_json'):
                            print("✅ Found b64_json attribute")
                            image_data = first_item.b64_json
                        elif hasattr(first_item, 'url'):
                            print("✅ Found url attribute") 
                            image_data = first_item.url
                        elif hasattr(first_item, 'image'):
                            print("✅ Found image attribute")
                            if isinstance(first_item.image, list) and len(first_item.image) > 0:
                                image_data = first_item.image[0]
                            else:
                                image_data = first_item.image
                        else:
                            print("❌ No recognized image data attributes found")
                            print(f"Available attributes: {[attr for attr in dir(first_item) if not attr.startswith('_')]}")
                            continue
                        
                        if image_data:
                            # Save image to local file
                            local_dir = "./generated_documents/images"
                            os.makedirs(local_dir, exist_ok=True)
                            
                            image_filename = f"image_{int(time.time())}_{i+1}.png"
                            image_path = os.path.join(local_dir, image_filename)
                            
                            print(f"💾 Saving image data type: {type(image_data)}")
                            print(f"💾 Image data preview: {str(image_data)[:100]}...")
                            
                            # Save the image data with improved handling
                            try:
                                if isinstance(image_data, str):
                                    if image_data.startswith('http'):
                                        # It's a URL, download the image
                                        print("🌐 Downloading image from URL...")
                                        import requests
                                        response = requests.get(image_data)
                                        with open(image_path, 'wb') as f:
                                            f.write(response.content)
                                    else:
                                        # Base64 encoded
                                        print("🔤 Decoding base64 image data...")
                                        import base64
                                        # Remove data URL prefix if present
                                        if ',' in image_data:
                                            image_data = image_data.split(',', 1)[1]
                                        with open(image_path, 'wb') as f:
                                            f.write(base64.b64decode(image_data))
                                else:
                                    # Binary data
                                    print("📦 Writing binary image data...")
                                    with open(image_path, 'wb') as f:
                                        f.write(image_data)
                                
                                generated_images.append({
                                    'prompt': prompt,
                                    'path': image_path,
                                    'filename': image_filename
                                })
                                
                                print(f"✅ Image {i+1} saved to {image_path}")
                                
                            except Exception as save_error:
                                print(f"❌ Error saving image {i+1}: {save_error}")
                        else:
                            print(f"❌ No image data extracted for prompt {i+1}")
                        
                    else:
                        print(f"❌ No valid response received for prompt {i+1}")
                        if response:
                            print(f"Response has data: {hasattr(response, 'data')}")
                            if hasattr(response, 'data'):
                                print(f"Data length: {len(response.data) if response.data else 0}")
                        
                except Exception as e:
                    print(f"❌ Error generating image {i+1}: {str(e)}")
                    continue
                    
        except ImportError:
            print("❌ OpenAI library not available for image generation")
            return []
        except Exception as e:
            print(f"❌ Error setting up image generation client: {str(e)}")
            return []
                
        return generated_images
    
    def _extract_image_prompts_robustly(self, content):
        """Robustly extract 3 image prompts from LLM response regardless of format."""
        try:
            import json
            import re
            
            print(f"🔍 Raw content to parse: {content[:200]}...")
            
            # Strategy 1: Try to parse as JSON array
            try:
                # Look for JSON array pattern
                json_match = re.search(r'\[.*?\]', content, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                    prompts = json.loads(json_str)
                    if isinstance(prompts, list) and len(prompts) >= 3:
                        print("✅ Successfully parsed JSON array format")
                        return [str(p).strip() for p in prompts[:3]]
            except:
                pass
            
            # Strategy 2: Try to parse as JSON object
            try:
                parsed = json.loads(content)
                if isinstance(parsed, dict):
                    # Look for various possible keys
                    possible_keys = [
                        ["Image1 Prompt", "Image2 Prompt", "Image3 Prompt"],
                        ["prompt1", "prompt2", "prompt3"],
                        ["1", "2", "3"],
                        ["image1", "image2", "image3"]
                    ]
                    
                    for key_set in possible_keys:
                        if all(key in parsed for key in key_set):
                            prompts = [parsed[key] for key in key_set]
                            if all(prompts):
                                print(f"✅ Successfully parsed JSON object with keys {key_set}")
                                return [str(p).strip() for p in prompts]
                
                # If it's a list directly
                elif isinstance(parsed, list) and len(parsed) >= 3:
                    print("✅ Successfully parsed direct JSON list")
                    return [str(p).strip() for p in parsed[:3]]
            except:
                pass
            
            # Strategy 3: Parse numbered lines (1. 2. 3. or 1) 2) 3))
            numbered_pattern = r'(?:^|\n)\s*(?:[1-3][\.\)])\s*(.+?)(?=(?:\n\s*[1-3][\.\)])|$)'
            numbered_matches = re.findall(numbered_pattern, content, re.MULTILINE | re.DOTALL)
            if len(numbered_matches) >= 3:
                prompts = [match.strip().replace('\n', ' ') for match in numbered_matches[:3]]
                if all(len(p) > 10 for p in prompts):  # Basic quality check
                    print("✅ Successfully parsed numbered list format")
                    return prompts
            
            # Strategy 4: Parse bullet points (- or • or *)
            bullet_pattern = r'(?:^|\n)\s*[-•*]\s*(.+?)(?=(?:\n\s*[-•*])|$)'
            bullet_matches = re.findall(bullet_pattern, content, re.MULTILINE | re.DOTALL)
            if len(bullet_matches) >= 3:
                prompts = [match.strip().replace('\n', ' ') for match in bullet_matches[:3]]
                if all(len(p) > 10 for p in prompts):  # Basic quality check
                    print("✅ Successfully parsed bullet point format")
                    return prompts
            
            # Strategy 5: Split by common delimiters and take first 3 substantial lines
            lines = []
            for delimiter in ['\n\n', '\n', '. ', '; ']:
                potential_lines = [line.strip() for line in content.split(delimiter)]
                substantial_lines = [line for line in potential_lines if len(line) > 20 and not line.lower().startswith(('based on', 'here are', 'the document', 'generate', 'create'))]
                if len(substantial_lines) >= 3:
                    lines = substantial_lines
                    break
            
            if len(lines) >= 3:
                print("✅ Successfully parsed by splitting and filtering lines")
                return lines[:3]
            
            # Strategy 6: Extract quoted strings
            quote_pattern = r'"([^"]{20,})"'
            quoted_matches = re.findall(quote_pattern, content)
            if len(quoted_matches) >= 3:
                print("✅ Successfully extracted quoted strings")
                return quoted_matches[:3]
            
            print("⚠️ Could not parse prompts from LLM response using any strategy")
            return None
            
        except Exception as e:
            print(f"❌ Error in robust parsing: {e}")
            return None
    
    def _create_intelligent_caption(self, image_prompt, figure_number):
        """Create an intelligent caption based on the image prompt."""
        try:
            # Extract key elements from the prompt to create a natural caption
            prompt_lower = image_prompt.lower()
            
            # Common themes and their caption templates
            if 'office' in prompt_lower and ('meeting' in prompt_lower or 'collaboration' in prompt_lower):
                return f"Figure {figure_number}: Professional team collaboration in a modern office environment"
            elif 'office' in prompt_lower and 'setting' in prompt_lower:
                return f"Figure {figure_number}: Contemporary office workspace and business environment" 
            elif 'workplace' in prompt_lower and 'diverse' in prompt_lower:
                return f"Figure {figure_number}: Diverse team members working together in a collaborative setting"
            elif 'meeting' in prompt_lower and ('room' in prompt_lower or 'conference' in prompt_lower):
                return f"Figure {figure_number}: Corporate meeting room with professional presentation setup"
            elif 'policy' in prompt_lower and 'document' in prompt_lower:
                return f"Figure {figure_number}: Professional documentation and policy review process"
            elif 'business' in prompt_lower and 'professional' in prompt_lower:
                return f"Figure {figure_number}: Professional business environment and corporate atmosphere"
            elif 'infographic' in prompt_lower or 'structure' in prompt_lower:
                return f"Figure {figure_number}: Visual representation of organizational structure and processes"
            elif 'corporate' in prompt_lower:
                return f"Figure {figure_number}: Corporate business environment and professional workspace"
            else:
                # Fallback: create a generic but descriptive caption
                # Extract first few meaningful words
                words = image_prompt.split()
                key_words = [w for w in words[:6] if w.lower() not in ['a', 'an', 'the', 'of', 'in', 'at', 'to', 'for', 'with']]
                if len(key_words) >= 3:
                    description = ' '.join(key_words[:4]).lower()
                    return f"Figure {figure_number}: Illustration depicting {description}"
                else:
                    return f"Figure {figure_number}: Professional business illustration"
                    
        except Exception as e:
            print(f"Warning: Error creating caption: {e}")
            return f"Figure {figure_number}: Professional business illustration"
    
    def _create_pdf_with_images(self, content, filename, doc_type, generated_images):
        """Create a PDF file with generated images embedded contextually."""
        
        # Ensure the volume directory exists (create locally for now)
        local_dir = "./generated_documents"
        os.makedirs(local_dir, exist_ok=True)
        
        pdf_path = os.path.join(local_dir, filename)
        
        # Create PDF with better margins and layout
        doc = SimpleDocTemplate(
            pdf_path, 
            pagesize=letter,
            leftMargin=1*inch,
            rightMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        styles = getSampleStyleSheet()
        
        # Enhanced custom styles for professional appearance
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=36,
            spaceBefore=12,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold',
            textColor='#2E3440'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=18,
            spaceBefore=24,
            alignment=0,  # Left alignment
            fontName='Helvetica-Bold',
            textColor='#3B4252'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=14,
            spaceBefore=2,
            alignment=4,  # Justified alignment
            fontName='Helvetica',
            leftIndent=0,
            rightIndent=0,
            firstLineIndent=0,
            leading=16,  # Line spacing
            textColor='#2E3440'
        )
        
        # Special style for bullet points and lists
        list_style = ParagraphStyle(
            'CustomList',
            parent=body_style,
            fontSize=12,
            spaceAfter=8,
            spaceBefore=4,
            leftIndent=20,
            bulletIndent=10,
            leading=15
        )
        
        image_caption_style = ParagraphStyle(
            'ImageCaption',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=16,
            spaceBefore=8,
            alignment=1,  # Center alignment
            fontName='Helvetica-Oblique',
            textColor='#5E81AC'
        )
        
        # Build document
        story = []
        
        # Add document header with company info and date
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")
        
        header_info = f"Generated on {current_date}"
        header_style = ParagraphStyle(
            'HeaderInfo',
            parent=styles['Normal'],
            fontSize=9,
            alignment=2,  # Right alignment
            textColor='#5E81AC',
            spaceAfter=20
        )
        story.append(Paragraph(header_info, header_style))
        
        # Title with better spacing
        title = f"{self._format_doc_type(doc_type)} Document"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Sanitize content for PDF generation
        content = self._sanitize_content_for_pdf(content)
        
        # Split content into paragraphs for image insertion
        paragraphs = content.split('\n\n')
        
        # Calculate where to insert images (distribute them evenly)
        total_paragraphs = len([p for p in paragraphs if p.strip()])
        if total_paragraphs > 0 and generated_images:
            # Insert images at 1/4, 1/2, and 3/4 through the document
            image_positions = [
                max(1, int(total_paragraphs * 0.25)),
                max(2, int(total_paragraphs * 0.5)),
                max(3, int(total_paragraphs * 0.75))
            ]
        else:
            image_positions = []
        
        paragraph_count = 0
        image_index = 0
        
        # Process each paragraph and insert images at calculated positions
        for para in paragraphs:
            if para.strip():
                try:
                    # Clean the paragraph text for ReportLab
                    para_text = para.strip()
                    
                    # Escape special characters that might cause ReportLab issues
                    para_text = para_text.replace('&', '&amp;')
                    para_text = para_text.replace('<', '&lt;')
                    para_text = para_text.replace('>', '&gt;')
                    
                    # Enhanced paragraph handling with better formatting detection
                    # Detect different paragraph types for appropriate styling
                    import re
                    
                    # Table of Contents
                    if 'TABLE OF CONTENTS' in para_text.upper():
                        toc_style = ParagraphStyle(
                            'TOC_Header',
                            parent=subtitle_style,
                            fontSize=16,
                            alignment=1,  # Center
                            spaceAfter=20,
                            spaceBefore=20
                        )
                        story.append(Paragraph(para_text, toc_style))
                    
                    # Numbered sections (1. 2. 3. etc.)
                    elif re.match(r'^\d+\.\s+[A-Z]', para_text):
                        story.append(Paragraph(para_text, subtitle_style))
                    
                    # Subsections (3.1, 4.2, etc.)  
                    elif re.match(r'^\s*\d+\.\d+\s+', para_text):
                        subsection_style = ParagraphStyle(
                            'Subsection',
                            parent=body_style,
                            fontSize=12,
                            leftIndent=20,
                            spaceBefore=8,
                            spaceAfter=8,
                            fontName='Helvetica-Bold'
                        )
                        story.append(Paragraph(para_text, subsection_style))
                    
                    # Horizontal rules and separators
                    elif '─' in para_text or '―' in para_text:
                        separator_style = ParagraphStyle(
                            'Separator',
                            parent=body_style,
                            alignment=1,  # Center
                            spaceBefore=12,
                            spaceAfter=12
                        )
                        story.append(Paragraph(para_text, separator_style))
                    
                    # Document metadata (Effective Date, Page X of Y, etc.)
                    elif any(keyword in para_text for keyword in ['Effective Date:', 'Document #', 'Page ', 'Revision #:']):
                        metadata_style = ParagraphStyle(
                            'Metadata',
                            parent=body_style,
                            fontSize=10,
                            textColor='#5E81AC',
                            alignment=1,  # Center
                            spaceBefore=8,
                            spaceAfter=16
                        )
                        story.append(Paragraph(para_text, metadata_style))
                    
                    # Section headers (ending with : or all caps)
                    elif len(para_text) < 80 and (para_text.endswith(':') or para_text.isupper()):
                        story.append(Paragraph(para_text, subtitle_style))
                    
                    # Bullet points and lists
                    elif para_text.startswith('•') or para_text.startswith('-') or para_text.startswith('*') or para_text.strip().startswith('•'):
                        story.append(Paragraph(para_text, list_style))
                    
                    # Policy/procedure headers 
                    elif len(para_text) < 120 and any(word in para_text.lower() for word in ['policy', 'procedure', 'guideline', 'section', 'article', 'roles', 'responsibilities', 'compliance']):
                        story.append(Paragraph(para_text, subtitle_style))
                    
                    # Company name and document titles (likely bold text)
                    elif len(para_text) < 100 and any(word in para_text for word in ['Inc.', 'LLC', 'Corp', 'Company', 'Solutions', 'Policy Guide', 'Human Resources']):
                        title_emphasis_style = ParagraphStyle(
                            'TitleEmphasis',
                            parent=body_style,
                            fontSize=13,
                            fontName='Helvetica-Bold',
                            alignment=1,  # Center
                            spaceBefore=12,
                            spaceAfter=12
                        )
                        story.append(Paragraph(para_text, title_emphasis_style))
                    
                    # Regular body text
                    else:
                        story.append(Paragraph(para_text, body_style))
                    
                    # Smart conditional spacing based on content type
                    if 'TABLE OF CONTENTS' in para_text.upper():
                        story.append(Spacer(1, 0.2*inch))
                    elif re.match(r'^\d+\.\s+[A-Z]', para_text) or len(para_text) < 80 and para_text.endswith(':'):
                        story.append(Spacer(1, 0.15*inch))  # More space after headers
                    elif '─' in para_text or any(keyword in para_text for keyword in ['Effective Date:', 'Document #']):
                        story.append(Spacer(1, 0.12*inch))  # Medium space for metadata
                    else:
                        story.append(Spacer(1, 0.08*inch))  # Standard spacing
                    
                    paragraph_count += 1
                    
                    # Check if we should insert an image after this paragraph
                    if (image_index < len(generated_images) and 
                        image_index < len(image_positions) and 
                        paragraph_count >= image_positions[image_index]):
                        
                        image_info = generated_images[image_index]
                        image_path = image_info['path']
                        
                        if os.path.exists(image_path):
                            try:
                                # Add some space before image
                                story.append(Spacer(1, 0.2*inch))
                                
                                # Create image with appropriate sizing
                                # Calculate size to fit within page margins (6 inches width max)
                                img = Image(image_path, width=5*inch, height=3*inch)
                                story.append(img)
                                
                                # Add intelligent image caption based on the prompt
                                image_prompt = image_info['prompt']
                                caption = self._create_intelligent_caption(image_prompt, image_index + 1)
                                story.append(Paragraph(caption, image_caption_style))
                                
                                # Add space after image
                                story.append(Spacer(1, 0.2*inch))
                                
                                print(f"📷 Inserted image {image_index + 1} after paragraph {paragraph_count}")
                                
                            except Exception as e:
                                print(f"Warning: Could not insert image {image_index + 1}: {str(e)}")
                        else:
                            print(f"Warning: Image file not found: {image_path}")
                        
                        image_index += 1
                        
                except Exception as e:
                    # If paragraph creation fails, add as plain text
                    print(f"Warning: Could not create paragraph, adding as plain text: {str(e)}")
                    # Convert to safe text and add as simple paragraph
                    safe_text = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    try:
                        story.append(Paragraph(safe_text, body_style))
                        story.append(Spacer(1, 0.1*inch))
                        paragraph_count += 1
                    except:
                        # Last resort: skip this paragraph
                        print(f"Skipping problematic paragraph: {para[:100]}...")
                        continue
        
        # Insert any remaining images at the end
        while image_index < len(generated_images):
            image_info = generated_images[image_index]
            image_path = image_info['path']
            
            if os.path.exists(image_path):
                try:
                    story.append(Spacer(1, 0.2*inch))
                    img = Image(image_path, width=5*inch, height=3*inch)
                    story.append(img)
                    # Add intelligent caption for end-of-document images too
                    image_prompt = image_info['prompt']
                    caption = self._create_intelligent_caption(image_prompt, image_index + 1)
                    story.append(Paragraph(caption, image_caption_style))
                    story.append(Spacer(1, 0.2*inch))
                    print(f"📷 Appended remaining image {image_index + 1} at document end")
                except Exception as e:
                    print(f"Warning: Could not insert final image {image_index + 1}: {str(e)}")
            
            image_index += 1
        
        # Build PDF
        doc.build(story)
        
        return pdf_path
    
    def _create_txt(self, content, filepath, company_name):
        """Create a text file with the generated content."""
        try:
            # Clean the content for text format
            clean_content = self._sanitize_content_for_text(content)
            
            # Add header information
            header = f"{company_name} - Generated Document\n"
            header += "=" * len(header.strip()) + "\n\n"
            header += f"Generated on: {datetime.now().strftime('%Y-%m-%d at %I:%M %p')}\n\n"
            
            # Combine header and content
            full_content = header + clean_content
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_content)
                
        except Exception as e:
            print(f"Error creating text file: {str(e)}")
            raise
    
    def _create_docx(self, content, filepath, company_name):
        """Create a Word document with the generated content."""
        try:
            # Clean the content for docx format
            clean_content = self._sanitize_content_for_text(content)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'{company_name} - Generated Document', 0)
            
            # Add generation timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_para.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d at %I:%M %p")}').italic = True
            
            # Add a line break
            doc.add_paragraph()
            
            # Process content - split by paragraphs and handle formatting
            paragraphs = clean_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a heading (starts with #, all caps, or ends with :)
                    if (para_text.strip().startswith('#') or 
                        para_text.strip().isupper() or 
                        para_text.strip().endswith(':')):
                        # Add as heading
                        heading_text = para_text.strip().lstrip('#').strip()
                        doc.add_heading(heading_text, level=1)
                    else:
                        # Add as regular paragraph
                        para = doc.add_paragraph()
                        # Handle basic formatting
                        lines = para_text.strip().split('\n')
                        for i, line in enumerate(lines):
                            if i > 0:
                                para.add_run('\n')
                            para.add_run(line.strip())
            
            # Save document
            doc.save(filepath)
            
        except Exception as e:
            print(f"Error creating Word document: {str(e)}")
            raise
    
    def _sanitize_content_for_text(self, content):
        """Sanitize content for text/docx format."""
        if not isinstance(content, str):
            content = str(content)
        
        # Remove HTML/XML tags
        import re
        content = re.sub(r'<[^>]+>', '', content)
        
        # Clean up excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        
        # Remove markdown table formatting for text
        lines = content.split('\n')
        clean_lines = []
        for line in lines:
            # Skip markdown table separator lines
            if re.match(r'^\s*\|[\s\-\|]*\|\s*$', line):
                continue
            # Convert table rows to simple text
            if '|' in line and not line.strip().startswith('|'):
                # This might be a table row, convert to simple format
                parts = [part.strip() for part in line.split('|') if part.strip()]
                if len(parts) > 1:
                    line = ' - '.join(parts)
            clean_lines.append(line)
        
        content = '\n'.join(clean_lines)
        
        # Ensure content doesn't start with newlines
        content = content.strip()
        
        return content

    def _save_to_volume(self, local_path, filename):
        """Save file to Databricks volume using WorkspaceClient pattern from cookbook."""
        import io
        from databricks.sdk import WorkspaceClient
        
        try:
            w = WorkspaceClient()
            
            # Read file into bytes
            print(f"Reading file: {local_path}")
            with open(local_path, "rb") as f:
                file_bytes = f.read()
            binary_data = io.BytesIO(file_bytes)
            
            # Construct volume file path
            # self.volume_path is "/Volumes/conor_smith/synthetic_data_app/synthetic_data_volume"
            volume_file_path = f"{self.volume_path}/{filename}"
            
            print(f"Uploading to volume: {volume_file_path}")
            w.files.upload(volume_file_path, binary_data, overwrite=True)
            
            print(f"✅ SUCCESS: {filename} uploaded to volume ({len(file_bytes):,} bytes)")
            return True
            
        except ImportError as e:
            print(f"❌ ERROR: Databricks SDK not available - {str(e)}")
            print("Install with: pip install databricks-sdk")
            return False
        except Exception as e:
            print(f"❌ ERROR: Upload failed - {str(e)}")
            print(f"   Source: {local_path}")
            print(f"   Target: {self.volume_path}/{filename}")
            return False

    def _list_volume_files(self):
        """List all files in the volume."""
        import io
        from databricks.sdk import WorkspaceClient
        
        try:
            w = WorkspaceClient()
            
            # List all files in the volume
            print(f"📋 Listing files in volume: {self.volume_path}")
            files = w.files.list_directory_contents(self.volume_path)
            
            file_list = []
            for file_info in files:
                if not file_info.is_directory:
                    file_list.append({
                        'name': file_info.name,
                        'path': file_info.path,
                        'size': file_info.file_size
                    })
                    
            print(f"📁 Found {len(file_list)} file(s) in volume")
            return file_list
            
        except ImportError as e:
            print(f"❌ ERROR: Databricks SDK not available - {str(e)}")
            return []
        except Exception as e:
            print(f"❌ ERROR: Failed to list volume files - {str(e)}")
            return []

    def _create_zip_from_local_files(self, timestamp):
        """Download missing files to local, then create zip from local files."""
        import zipfile
        import os
        from databricks.sdk import WorkspaceClient
        
        try:
            print(f"⚡ Smart download approach: ensure all files are local first...")
            
            # Get list of completed files from generation state
            completed_items = self.generation_state.get('completed_items', [])
            filenames = []
            
            # Extract filenames from completed items
            for item in completed_items:
                if isinstance(item, dict) and 'filename' in item:
                    filenames.append(item['filename'])
                elif isinstance(item, list):
                    for sub_item in item:
                        if isinstance(sub_item, dict) and 'filename' in sub_item:
                            filenames.append(sub_item['filename'])
            
            if not filenames:
                print("⚠️  No files found from current batch")
                return None
            
            # Ensure local directory exists
            local_dir = "./generated_documents"
            os.makedirs(local_dir, exist_ok=True)
            
            print(f"📋 Processing {len(filenames)} file(s) for download...")
            
            # Check which files exist locally and which need downloading
            available_files = []
            files_to_download = []
            
            for filename in filenames:
                local_path = os.path.join(local_dir, filename)
                if os.path.exists(local_path):
                    file_size = os.path.getsize(local_path)
                    available_files.append((filename, local_path, file_size))
                    print(f"   ✅ Already local: {filename} ({file_size:,} bytes)")
                else:
                    files_to_download.append(filename)
                    print(f"   📥 Need to download: {filename}")
            
            # Download missing files from volume
            if files_to_download:
                print(f"📡 Downloading {len(files_to_download)} file(s) from volume...")
                
                try:
                    w = WorkspaceClient()
                    
                    for i, filename in enumerate(files_to_download):
                        try:
                            print(f"   📥 [{i+1}/{len(files_to_download)}] Downloading: {filename}")
                            
                            # Download from volume using Databricks SDK
                            download_file_path = f"{self.volume_path}/{filename}"
                            print(f"      Volume path: {download_file_path}")
                            
                            response = w.files.download(download_file_path)
                            
                            # Handle StreamingResponse properly
                            if hasattr(response, 'contents'):
                                # Try reading from contents if it's a stream
                                if hasattr(response.contents, 'read'):
                                    file_data = response.contents.read()
                                else:
                                    file_data = response.contents
                            else:
                                # Fallback: try to read response directly
                                file_data = response.read() if hasattr(response, 'read') else response
                            
                            print(f"      Response type: {type(response)}")
                            print(f"      Contents type: {type(response.contents) if hasattr(response, 'contents') else 'No contents attr'}")
                            print(f"      File data type: {type(file_data)}")
                            
                            # Save to local directory
                            local_path = os.path.join(local_dir, filename)
                            with open(local_path, 'wb') as f:
                                f.write(file_data)
                            
                            file_size = len(file_data)
                            available_files.append((filename, local_path, file_size))
                            print(f"      ✅ Downloaded and saved locally: {file_size:,} bytes")
                            
                        except Exception as download_error:
                            print(f"      ❌ Failed to download {filename}: {str(download_error)}")
                            continue
                            
                except Exception as sdk_error:
                    print(f"❌ ERROR: Databricks SDK issue - {str(sdk_error)}")
                    if not available_files:
                        return None
            
            if not available_files:
                print("❌ No files available for zipping")
                return None
            
            # Create zip from all local files
            zip_filename = f"synthetic_data_batch_{timestamp}.zip"
            temp_zip_path = os.path.join(local_dir, zip_filename)
            
            print(f"📦 Creating zip from {len(available_files)} local file(s)...")
            
            successful_files = 0
            with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for filename, local_path, file_size in available_files:
                    try:
                        zipf.write(local_path, filename)
                        successful_files += 1
                        print(f"   ✅ Added to zip: {filename} ({file_size:,} bytes)")
                    except Exception as zip_error:
                        print(f"   ❌ Failed to add {filename}: {str(zip_error)}")
            
            if successful_files == 0:
                print("❌ No files were successfully added to zip")
                return None
            
            # Check final zip size
            zip_size = os.path.getsize(temp_zip_path) if os.path.exists(temp_zip_path) else 0
            print(f"🎯 Zip created successfully: {zip_size:,} bytes with {successful_files} files")
            
            return {
                'local_path': temp_zip_path,
                'filename': zip_filename,
                'saved_to_volume': False,  # Skip volume save for speed
                'file_count': successful_files
            }
            
        except Exception as e:
            print(f"❌ ERROR in smart download zip creation: {str(e)}")
            import traceback
            print(f"Full traceback: {traceback.format_exc()}")
            return None

    def _download_and_zip_files(self, timestamp):
        """Download all files from volume and create a zip file."""
        import zipfile
        import os
        import io
        import tempfile
        from databricks.sdk import WorkspaceClient
        
        try:
            print(f"🏁 Starting zip creation process...")
            w = WorkspaceClient()
            
            # Get list of completed files from generation state
            completed_items = self.generation_state.get('completed_items', [])
            filenames = []
            
            print(f"🔍 Analyzing completed items: {len(completed_items)} total")
            
            # Extract filenames from completed items
            for i, item in enumerate(completed_items):
                print(f"   Item {i+1}: {type(item)} - {item if isinstance(item, str) else 'dict/list'}")
                if isinstance(item, dict) and 'filename' in item:
                    filenames.append(item['filename'])
                    print(f"     → Found filename: {item['filename']}")
                elif isinstance(item, list):
                    # Handle batch operations that return lists
                    for j, sub_item in enumerate(item):
                        if isinstance(sub_item, dict) and 'filename' in sub_item:
                            filenames.append(sub_item['filename'])
                            print(f"     → Found filename in subitem {j+1}: {sub_item['filename']}")
            
            if not filenames:
                print("⚠️  No files found to zip from current batch generation")
                return None
            
            # Create zip file locally first  
            zip_filename = f"synthetic_data_batch_{timestamp}.zip"
            temp_zip_path = f"./generated_documents/{zip_filename}"
            os.makedirs('./generated_documents', exist_ok=True)
            
            print(f"📦 Creating zip file: {zip_filename}")
            print(f"   Files to include ({len(filenames)}):")
            for filename in filenames:
                print(f"     • {filename}")
            
            successful_files = 0
            
            with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for i, filename in enumerate(filenames):
                    try:
                        print(f"   📥 [{i+1}/{len(filenames)}] Processing: {filename}")
                        
                        # Download file from volume
                        volume_file_path = f"{self.volume_path}/{filename}"
                        print(f"      Volume path: {volume_file_path}")
                        
                        # Use Databricks SDK to download file
                        print(f"      Downloading from volume...")
                        file_content = w.files.download(volume_file_path)
                        
                        print(f"      Downloaded {len(file_content.contents):,} bytes")
                        
                        # Add to zip
                        zipf.writestr(filename, file_content.contents)
                        successful_files += 1
                        print(f"      ✅ Added to zip successfully")
                        
                    except Exception as file_error:
                        print(f"      ❌ Failed to process {filename}: {str(file_error)}")
                        continue
            
            print(f"🎯 Zip creation complete: {successful_files}/{len(filenames)} files added successfully")
            
            if successful_files == 0:
                print("❌ No files were successfully added to zip")
                return None
            
            # Check zip file size
            zip_size = os.path.getsize(temp_zip_path) if os.path.exists(temp_zip_path) else 0
            print(f"📊 Created zip file: {zip_size:,} bytes")
            
            # Skip saving zip to volume to avoid timeouts
            print(f"⚡ Skipping volume save for faster download...")
            zip_saved_to_volume = False
            
            return {
                'local_path': temp_zip_path,
                'filename': zip_filename,
                'saved_to_volume': zip_saved_to_volume,
                'file_count': successful_files
            }
            
        except ImportError as e:
            print(f"❌ ERROR: Databricks SDK not available - {str(e)}")
            return None
        except Exception as e:
            print(f"❌ ERROR: Failed to create zip file - {str(e)}")
            import traceback
            print(f"Full traceback: {traceback.format_exc()}")
            return None

    def _format_doc_type(self, doc_type):
        """Format document type for display."""
        type_map = {
            'policy_guide': 'Policy Guide',
            'customer_correspondence': 'Customer Correspondence',
            'customer_profile': 'Customer Profile',
            'custom_document': 'Custom Document'
        }
        return type_map.get(doc_type, doc_type.replace('_', ' ').title())

    def _sanitize_content_for_pdf(self, content):
        """Enhanced content cleaning for proper PDF formatting with better line breaks and structure."""
        if not isinstance(content, str):
            content = str(content)
        
        import re
        
        # Remove HTML/XML tags
        content = re.sub(r'<[^>]+>', '', content)
        
        # Step 1: Handle bold/italic markdown formatting
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # **bold** → bold
        content = re.sub(r'\*(.*?)\*', r'\1', content)      # *italic* → italic
        content = re.sub(r'__(.*?)__', r'\1', content)      # __bold__ → bold  
        content = re.sub(r'_(.*?)_', r'\1', content)        # _italic_ → italic
        
        # Step 2: Add proper line breaks around major sections and headers
        # Handle document headers and separators
        content = re.sub(r'(\*Document #.*?\*)', r'\n\n\1\n', content)
        content = re.sub(r'(Effective Date:.*?)([A-Z])', r'\1\n\n\2', content)
        content = re.sub(r'(Page \d+ of \d+)', r'\1\n', content)
        content = re.sub(r'---+', '\n' + '─' * 60 + '\n', content)  # Pretty horizontal rules
        
        # Step 3: Improve Table of Contents formatting
        content = re.sub(r'### Table of Contents', '\n\nTABLE OF CONTENTS\n', content)
        content = re.sub(r'## Table of Contents', '\n\nTABLE OF CONTENTS\n', content)
        content = re.sub(r'# Table of Contents', '\n\nTABLE OF CONTENTS\n', content)
        
        # Step 4: Handle numbered sections properly - add line breaks before major section numbers
        content = re.sub(r'(\d+\.\s*\*\*[^*]+\*\*)', r'\n\n\1', content)  # 1. **Section Title**
        content = re.sub(r'(\d+\.\s*[A-Z][^.]*[a-z])', r'\n\n\1', content)  # 1. Section Title
        
        # Step 5: Format subsections with proper indentation
        content = re.sub(r'(\d+\.\d+\s+[A-Za-z])', r'\n    \1', content)  # 3.1 Subsection
        
        # Step 6: Handle dotted lines in table of contents
        content = re.sub(r'\.{3,}', ' ' + '·' * 20 + ' ', content)  # Dotted leaders
        
        # Step 7: Add line breaks around list items and improve spacing
        content = re.sub(r'(\d+\.\s+[A-Z][^0-9\n]*?)(\s+\d+\.)', r'\1\n\n\2', content)
        
        # Step 8: Handle special sections like roles, compliance, etc.
        special_sections = [
            r'(Roles & Responsibilities)',
            r'(Compliance & Reporting)', 
            r'(Special Circumstances)',
            r'(Policy Statement)',
            r'(Purpose and Scope)',
            r'(Request and Approval Process)',
            r'(Use of Vacation Time)'
        ]
        for pattern in special_sections:
            content = re.sub(pattern, r'\n\n\1', content)
        
        # Step 9: Convert markdown tables to readable text
        lines = content.split('\n')
        cleaned_lines = []
        in_table = False
        
        for line in lines:
            # Detect table rows (lines with multiple | characters)
            if '|' in line and line.count('|') >= 2:
                in_table = True
                # Clean up table formatting
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:  # Skip empty rows
                    # Skip header separator rows (like |------|-----|)
                    if not all(cell.replace('-', '').replace(' ', '') == '' for cell in cells):
                        # Format as readable text
                        if len(cells) >= 2:
                            cleaned_lines.append(f"  • {cells[0]}: {' '.join(cells[1:])}")
                        else:
                            cleaned_lines.append(f"  • {' '.join(cells)}")
            else:
                if in_table and line.strip() == '':
                    in_table = False
                    cleaned_lines.append('')  # Add space after table
                elif not in_table:
                    cleaned_lines.append(line)
        
        content = '\n'.join(cleaned_lines)
        
        # Step 10: Clean up special characters that might cause issues
        content = content.replace('\u2022', '•')  # Bullet points
        content = content.replace('\u2011', '-')  # Non-breaking hyphens
        content = content.replace('\u2013', '-')  # En dashes
        content = content.replace('\u2014', '—') # Em dashes
        content = content.replace('\u201c', '"')  # Left double quote
        content = content.replace('\u201d', '"')  # Right double quote
        content = content.replace('\u2018', "'")  # Left single quote
        content = content.replace('\u2019', "'")  # Right single quote
        
        # Step 11: Remove any remaining HTML entities
        content = re.sub(r'&[a-zA-Z0-9#]+;', '', content)
        
        # Step 12: Clean up excessive whitespace while preserving intentional spacing
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Multiple blank lines → double
        content = re.sub(r'[ \t]+', ' ', content)  # Multiple spaces/tabs → single space
        content = re.sub(r' +\n', '\n', content)  # Trailing spaces
        
        # Step 13: Ensure proper spacing around important elements
        content = re.sub(r'(TABLE OF CONTENTS)', r'\n\1\n', content)
        content = re.sub(r'([.!?])\s*([A-Z][a-z])', r'\1\n\n\2', content)  # Sentence breaks
        
        return content.strip()

    def _extract_content_safely(self, response):
        """Safely extract content from various response formats."""
        if isinstance(response, dict):
            # Handle reasoning/summary structure first
            if response.get('type') == 'reasoning' and 'summary' in response:
                # This is metadata, not the actual content - skip it
                print("Detected reasoning metadata, looking for actual content...")
                return "Content extraction failed - received metadata instead of document content"
            
            # Try common keys in order of preference
            for key in ["content", "text", "message", "summary"]:
                if key in response:
                    value = response[key]
                    # Handle nested structures
                    if isinstance(value, dict):
                        if "content" in value:
                            return value["content"]
                        elif "text" in value:
                            return value["text"]
                        else:
                            # Skip metadata objects
                            continue
                    elif isinstance(value, list):
                        # If it's a list, try to extract content from each item
                        extracted_items = []
                        for item in value:
                            if isinstance(item, dict):
                                # Skip metadata objects like reasoning summaries
                                if item.get('type') in ['summary_text', 'reasoning']:
                                    continue
                                elif "content" in item:
                                    extracted_items.append(item["content"])
                                elif "text" in item:
                                    extracted_items.append(item["text"])
                                else:
                                    # Only include if it looks like actual content
                                    item_str = str(item)
                                    if len(item_str) > 100 and not item_str.startswith("{'type'"):
                                        extracted_items.append(item_str)
                            else:
                                extracted_items.append(str(item))
                        return '\n\n'.join(extracted_items) if extracted_items else str(value)
                    else:
                        # Only return string values that look like actual content
                        value_str = str(value)
                        if len(value_str) > 50 and not value_str.startswith("{'type'"):
                            return value_str
            
            # If none of the expected keys exist, return string representation
            response_str = str(response)
            # But avoid returning obvious metadata
            if response_str.startswith("{'type': 'reasoning'"):
                return "Error: Received reasoning metadata instead of document content"
            return response_str
            
        elif isinstance(response, list):
            # If response is directly a list, extract content from each item
            extracted_items = []
            reasoning_content = None
            
            for item in response:
                if isinstance(item, dict):
                    # Look for actual text content first
                    if item.get('type') == 'text' and 'text' in item:
                        extracted_items.append(item['text'])
                    elif item.get('type') == 'reasoning' and 'summary' in item:
                        # If we only have reasoning, try to extract useful content from it
                        summary = item['summary']
                        if isinstance(summary, list):
                            for summary_item in summary:
                                if isinstance(summary_item, dict) and 'text' in summary_item:
                                    summary_text = summary_item['text']
                                    # Look for quoted content that might be the intended response
                                    import re
                                    quotes = re.findall(r'"([^"]+)"', summary_text)
                                    for quote in quotes:
                                        if len(quote) > 30 and not quote.startswith("Give me"):
                                            reasoning_content = quote
                                            break
                    else:
                        # Try recursive extraction for other types
                        content = self._extract_content_safely(item)
                        if content and not content.startswith("Error:") and len(content) > 50:
                            extracted_items.append(content)
                else:
                    extracted_items.append(str(item))
            
            # If we found actual content, use that
            if extracted_items:
                return '\n\n'.join(extracted_items)
            # Otherwise, fall back to reasoning content if available
            elif reasoning_content:
                return reasoning_content
            else:
                return "Error: Unable to extract usable content from LLM response"
        else:
            return str(response)

    def _add_custom_css(self):
        """Add custom CSS for animations and styling."""
        custom_css = '''
        @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css');
        @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;700&display=swap');
        
        body {
            font-family: 'DM Sans', sans-serif;
            background-color: #F9F7F4;
        }
        
        .fa-spinner {
            animation: spin 1s linear infinite;
        }
        
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        .generate-button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
        }
        
        .progress-container {
            border-radius: 10px;
            overflow: hidden;
        }
        
        .alert-info {
            border-left: 4px solid #17a2b8;
        }
        
        .alert-success {
            border-left: 4px solid #28a745;
        }
        
        .alert-danger {
            border-left: 4px solid #dc3545;
        }
        
        .document-config-card {
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            border-radius: 10px;
        }
        
        .form-label {
            color: #1B3139;
            font-weight: 600;
        }
        
        /* Countdown timer styles */
        .countdown-timer {
            font-family: 'Courier New', monospace;
            font-weight: bold;
            color: #007bff;
            font-size: 2rem;
        }
        '''
        
        self.app.index_string = self.app.index_string.replace(
            '</head>',
            f'<style>{custom_css}</style></head>'
        )
    
    def _format_data_type(self, data_type):
        """Format data type for display."""
        type_map = {
            'pdf': 'PDF Document',
            'text': 'Text Document',
            'tabular': 'Tabular Data'
        }
        return type_map.get(data_type, data_type.title())
    
    def _format_items_display(self, items, data_type):
        """Format items for history display."""
        if not items:
            return html.P("No items generated", className="text-muted small")
        
        item_elements = []
        for item in items:
            if data_type == 'pdf':
                icon = "📄"
                size_info = f" ({item.get('size', 0):,} bytes)" if item.get('size') else ""
                volume_status = " ✅" if item.get('saved_to_volume') else " 📁"
            elif data_type == 'text':
                icon = "📝"
                size_info = f" ({item.get('size', 0):,} chars)" if item.get('size') else ""
                volume_status = ""
            elif data_type == 'tabular':
                icon = "📊"
                size_info = f" ({item.get('size', 0):,} bytes)" if item.get('size') else ""
                volume_status = ""
            else:
                icon = "📄"
                size_info = ""
                volume_status = ""
            
            item_elements.append(
                html.Li([
                    f"{icon} {item.get('filename', 'Unknown file')}",
                    html.Small(size_info + volume_status, className="text-muted")
                ], className="small")
            )
        
        return html.Ul(item_elements, className="mb-0")
    
    def _generate_data_background(self, data_type, description, company_name, company_sector):
        """Generate data in background thread with progress updates."""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Update progress
            self.generation_state['current_item'] = 1
            self.generation_state['current_step'] = f"Generating {self._format_data_type(data_type)}..."
            
            if data_type == 'pdf':
                # Generate PDF (existing functionality)
                item_info = self._generate_pdf_item(description, company_name, company_sector, timestamp, 'policy_guide', False, '')
            elif data_type == 'text':
                # Generate text document
                item_info = self._generate_text_item(description, company_name, company_sector, timestamp)
            elif data_type == 'tabular':
                # Generate tabular data
                item_info = self._generate_tabular_item(description, company_name, company_sector, timestamp)
            else:
                raise Exception(f"Unknown data type: {data_type}")
            
            self.generation_state['completed_items'].append(item_info)
            
            # Mark as complete
            self.generation_state['active'] = False
            self.generation_state['current_step'] = 'Complete!'
            
        except Exception as e:
            print(f"Error in background generation: {str(e)}")
            self.generation_state['active'] = False
            self.generation_state['error'] = str(e)

    def _generate_pdf_item(self, description, company_name, company_sector, timestamp, doc_type='policy_guide', include_images=False, doc_name=''):
        """Generate a single PDF item with optional AI-generated images."""
        # Enhanced prompt with company context
        enhanced_description = f"For {company_name} (a {company_sector} company): {description}"
        
        # Generate content using the serving endpoint
        content = self._generate_document_content(doc_type, enhanced_description, 1)
        
        # Create filename using document name if provided, otherwise fallback to synthetic_pdf
        if doc_name:
            # Sanitize document name for filename
            safe_name = "".join(c for c in doc_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_name = safe_name.replace(' ', '_')
            filename = f"{safe_name}_{timestamp}.pdf"
        else:
            filename = f"synthetic_pdf_{timestamp}.pdf"
        
        # Agentic workflow for image generation
        generated_images = []
        if include_images:
            try:
                # Step 1: Generate image prompts based on PDF content
                self.generation_state['current_step'] = f"🤖 Analyzing content for image prompts..."
                print("🧠 AGENTIC FLOW: Generating image prompts from PDF content...")
                image_prompts = self._generate_image_prompts(content)
                print(f"📝 Generated {len(image_prompts)} image prompts")
                
                # Step 2: Generate images from prompts
                if image_prompts:
                    self.generation_state['current_step'] = f"🎨 Generating {len(image_prompts)} AI images..."
                    print("🎨 AGENTIC FLOW: Creating images from prompts...")
                    generated_images = self._generate_images(image_prompts)
                    print(f"✅ Generated {len(generated_images)} images successfully")
                    
            except Exception as e:
                print(f"⚠️ Error in image generation workflow: {str(e)}")
                print("🔄 Continuing with PDF generation without images...")
                generated_images = []
        
        # Update progress
        self.generation_state['current_step'] = f"📄 Creating PDF{' with images' if generated_images else ''}..."
        
        # Generate PDF with or without images
        if generated_images:
            print(f"📋 Creating PDF with {len(generated_images)} embedded images...")
            pdf_path = self._create_pdf_with_images(content, filename, doc_type, generated_images)
        else:
            print("📄 Creating standard PDF...")
            pdf_path = self._create_pdf(content, filename, doc_type)
        
        # Update progress and save to volume
        self.generation_state['current_step'] = f"💾 Saving PDF to volume..."
        
        try:
            self._save_to_volume(pdf_path, filename)
            save_success = True
        except Exception as e:
            save_success = False
            print(f"Failed to save PDF to volume: {str(e)}")
            # For iterative generation, we don't fail completely on volume save errors
        
        return {
            'type': 'pdf',
            'filename': filename,
            'path': pdf_path,
            'description': description,
            'company_name': company_name,
            'company_sector': company_sector,
            'size': os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0,
            'saved_to_volume': save_success,
            'include_images': include_images,
            'images_generated': len(generated_images)
        }

    def _generate_text_item(self, description, company_name, company_sector, timestamp, doc_type, file_format, doc_name=''):
        """Generate a single text document with LLM content."""
        try:
            # Generate content using LLM (similar to PDF generation)
            self.generation_state['current_step'] = f"Generating {file_format.upper()} content..."
            
            # Enhance description with company context
            enhanced_description = f"""For {company_name} (a {company_sector} company): {description}
            
Company Context:
- Company Name: {company_name}
- Industry Sector: {company_sector}

Please incorporate this company information naturally throughout the document to make it specific and realistic for this organization."""
            
            # Extract doc number from timestamp for document series numbering
            doc_number = int(timestamp.split('_')[-1]) if '_' in timestamp else 1
            content = self._generate_document_content(doc_type, enhanced_description, doc_number)
            
            # Create filename using document name if provided, otherwise fallback to doc_type
            if doc_name:
                # Sanitize document name for filename
                safe_name = "".join(c for c in doc_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                safe_name = safe_name.replace(' ', '_')
                if file_format == 'docx':
                    filename = f"{safe_name}_{timestamp}.docx"
                    extension = 'docx'
                else:
                    filename = f"{safe_name}_{timestamp}.txt"
                    extension = 'txt'
            else:
                if file_format == 'docx':
                    filename = f"{doc_type}_{timestamp}.docx"
                    extension = 'docx'
                else:
                    filename = f"{doc_type}_{timestamp}.txt"
                    extension = 'txt'
                
            # Create local file
            os.makedirs('./generated_documents', exist_ok=True)
            local_path = f"./generated_documents/{filename}"
            
            if file_format == 'docx':
                self._create_docx(content, local_path, company_name)
            else:
                self._create_txt(content, local_path, company_name)
            
            # Save to volume
            self.generation_state['current_step'] = f"Saving {extension.upper()} to volume..."
            volume_success = self._save_to_volume(local_path, filename)
            
            return {
                'type': 'text',
                'description': description,
                'filename': filename,
                'format': file_format,
                'doc_type': doc_type,
                'path': local_path,
                'timestamp': timestamp,
                'company_name': company_name,
                'company_sector': company_sector,
                'saved_to_volume': volume_success,
                'size': os.path.getsize(local_path) if os.path.exists(local_path) else 0
            }
            
        except Exception as e:
            print(f"Error generating text document: {str(e)}")
            return {
                'type': 'text',
                'description': description,
                'filename': f"error_{timestamp}.txt",
                'format': file_format,
                'error': str(e),
                'timestamp': timestamp,
                'company_name': company_name,
                'company_sector': company_sector,
                'saved_to_volume': False,
                'path': None,
                'size': 0
            }

    def _generate_tabular_item(self, table_name, row_count, columns, company_name, company_sector, timestamp):
        """Generate tabular data using Databricks job on remote cluster."""
        try:
            self.generation_state['current_step'] = f"Starting Databricks job for {row_count} rows of tabular data..."
            
            # Import Databricks SDK
            from databricks.sdk import WorkspaceClient
            import json
            
            print("🚀 TABULAR GENERATION: Triggering Databricks job")
            print(f"   - Table name: {table_name}")
            print(f"   - Row count: {row_count}")
            print(f"   - Column count: {len(columns)}")
            
            # Initialize Databricks client
            w = WorkspaceClient()
            
            # Job ID for the tabular data generation job
            job_id = "635184344270819"
            
            # Prepare parameters for the job
            job_parameters = {
                "table_name": table_name,
                "row_count": str(row_count),
                "columns": json.dumps(columns),  # Serialize column configurations
                "company_name": company_name,
                "company_sector": company_sector,
                "timestamp": timestamp,
                "endpoint_name": self.endpoint_name,
                "volume_path": "conor_smith/synthetic_data_app/synthetic_data_volume"
            }
            
            self.generation_state['current_step'] = f"Submitting job to Databricks cluster..."
            
            # Start the job
            try:
                print("🔄 DATABRICKS JOB: Submitting job with parameters:")
                print(f"   - Job ID: {job_id}")
                for key, value in job_parameters.items():
                    if key == "columns":
                        print(f"   - {key}: [Column configurations - {len(columns)} columns]")
                    else:
                        print(f"   - {key}: {value}")
                
                # Try to run the job with parameters
                print(f"📤 DATABRICKS JOB: Attempting to start job with parameters...")
                
                try:
                    run = w.jobs.run_now(job_id=job_id, job_parameters=job_parameters)
                    print(f"✅ DATABRICKS JOB: Started run with ID {run.run_id}")
                except Exception as job_start_error:
                    print(f"⚠️  DATABRICKS JOB: Error starting job with parameters: {job_start_error}")
                    print("🔄 DATABRICKS JOB: Attempting to start job without parameters as fallback...")
                    
                    # Try without parameters as fallback
                    try:
                        run = w.jobs.run_now(job_id=job_id)
                        print(f"✅ DATABRICKS JOB: Started run without parameters, ID: {run.run_id}")
                        print("⚠️  Note: Job may use default parameters")
                    except Exception as fallback_error:
                        print(f"❌ DATABRICKS JOB: Failed to start job even without parameters: {fallback_error}")
                        raise fallback_error
                
                self.generation_state['current_step'] = f"Job running on Databricks cluster (Run ID: {run.run_id})..."
                
                # Wait for job completion (with timeout)
                print(f"⏳ DATABRICKS JOB: Waiting for completion...")
                import time
                
                timeout_seconds = 600  # 10 minute timeout
                start_time = time.time()
                
                while True:
                    if time.time() - start_time > timeout_seconds:
                        print(f"⚠️  DATABRICKS JOB: Timeout after {timeout_seconds} seconds")
                        break
                    
                    try:
                        run_status = w.jobs.get_run(run.run_id)
                        state = run_status.state.life_cycle_state.value if run_status.state else "UNKNOWN"
                        
                        if state in ["TERMINATED", "SKIPPED", "INTERNAL_ERROR"]:
                            result_state = run_status.state.result_state.value if run_status.state and run_status.state.result_state else "UNKNOWN"
                            print(f"📋 DATABRICKS JOB: Completed with state: {state}, result: {result_state}")
                            
                            if result_state == "SUCCESS":
                                print("✅ DATABRICKS JOB: Tabular data generation completed successfully")
                                break
                            else:
                                print(f"❌ DATABRICKS JOB: Failed with result: {result_state}")
                                break
                        else:
                            print(f"🔄 DATABRICKS JOB: Status: {state}")
                            time.sleep(10)  # Wait 10 seconds before checking again
                            
                    except Exception as e:
                        print(f"⚠️  Error checking job status: {e}")
                        time.sleep(10)
                
                # Generate return data (job handles file creation and volume storage)
                filename = f"{table_name}_{timestamp}.csv"
                
                return {
                    'type': 'tabular',
                    'table_name': table_name,
                    'filename': filename,
                    'job_run_id': run.run_id,
                    'row_count': row_count,
                    'column_count': len(columns),
                    'columns': [col.get('name', 'unnamed') for col in columns],
                    'preview_data': "Generated by Databricks job - check volume for full data",
                    'timestamp': timestamp,
                    'company_name': company_name,
                    'company_sector': company_sector,
                    'size': 0,  # Size will be determined by the job
                    'saved_to_volume': True,  # Job handles volume storage
                    'generation_method': 'databricks_job'
                }
                
            except Exception as e:
                print(f"❌ DATABRICKS JOB: Failed to start job - {str(e)}")
                raise e
                
        except Exception as e:
            print(f"⚠️  TABULAR GENERATION: Databricks job failed ({str(e)})")
            print("🔄 FALLBACK: Using Pandas implementation")
            return self._generate_tabular_fallback(table_name, row_count, columns, company_name, company_sector, timestamp)

    def _generate_tabular_fallback(self, table_name, row_count, columns, company_name, company_sector, timestamp):
        """Fallback method to generate tabular data without Spark."""
        print("🔄 PANDAS FALLBACK: Starting fallback tabular generation")
        print(f"   - Row count: {row_count}")
        print(f"   - GenAI Text will generate unique content for ALL rows")
        
        try:
            import pandas as pd
            import random
            from faker import Faker
            fake = Faker()
            
            self.generation_state['current_step'] = f"Generating {row_count} rows (fallback mode)..."
            
            # Create data dictionary
            data = {}
            
            # Generate data for each column
            for col in columns:
                col_name = col.get('name', 'unnamed_column')
                col_type = col.get('data_type', 'Integer')
                
                if col_type == 'Integer':
                    min_val = col.get('min_value', 1)
                    max_val = col.get('max_value', 100)
                    data[col_name] = [random.randint(min_val, max_val) for _ in range(row_count)]
                elif col_type == 'Date':
                    # Generate random dates between min_date and max_date
                    from datetime import datetime, timedelta
                    min_date_str = col.get('min_date', '2020-01-01')
                    max_date_str = col.get('max_date', '2024-12-31')
                    
                    try:
                        min_date = datetime.strptime(min_date_str, '%Y-%m-%d')
                        max_date = datetime.strptime(max_date_str, '%Y-%m-%d')
                        date_range = (max_date - min_date).days
                        
                        # Generate random dates
                        random_dates = []
                        for _ in range(row_count):
                            random_days = random.randint(0, date_range)
                            random_date = min_date + timedelta(days=random_days)
                            random_dates.append(random_date.strftime('%Y-%m-%d'))
                        
                        data[col_name] = random_dates
                    except Exception as date_error:
                        print(f"⚠️  Error generating dates for {col_name}: {date_error}")
                        # Fallback to default dates
                        data[col_name] = ['2022-01-01'] * row_count
                elif col_type == 'First Name':
                    data[col_name] = [fake.first_name() for _ in range(row_count)]
                elif col_type == 'Last Name':
                    data[col_name] = [fake.last_name() for _ in range(row_count)]
                elif col_type == 'Email':
                    data[col_name] = [fake.email() for _ in range(row_count)]
                elif col_type == 'Phone Number':
                    data[col_name] = [fake.phone_number() for _ in range(row_count)]
                elif col_type == 'Address':
                    data[col_name] = [fake.address() for _ in range(row_count)]
                elif col_type == 'City':
                    data[col_name] = [fake.city() for _ in range(row_count)]
                elif col_type == 'Country':
                    data[col_name] = [fake.country() for _ in range(row_count)]
                elif col_type == 'Country Code':
                    data[col_name] = [fake.country_code() for _ in range(row_count)]
                elif col_type == 'Postcode':
                    data[col_name] = [fake.postcode() for _ in range(row_count)]
                elif col_type == 'Street Address':
                    data[col_name] = [fake.street_address() for _ in range(row_count)]
                elif col_type == 'Company':
                    data[col_name] = [fake.company() for _ in range(row_count)]
                elif col_type == 'Credit Card Number':
                    data[col_name] = [fake.credit_card_number() for _ in range(row_count)]
                elif col_type == 'Credit Card Provider':
                    data[col_name] = [fake.credit_card_provider() for _ in range(row_count)]
                elif col_type == 'Latitude':
                    # Convert decimal.Decimal to float for proper numeric handling
                    data[col_name] = [float(fake.latitude()) for _ in range(row_count)]
                elif col_type == 'Longitude':
                    # Convert decimal.Decimal to float for proper numeric handling
                    data[col_name] = [float(fake.longitude()) for _ in range(row_count)]
                elif col_type == 'Country Defined Coordinates':
                    # Generate location data for specific country
                    country_code = col.get('country_code', 'US')
                    coords_only = col.get('coords_only', False)
                    
                    # Generate location tuples
                    locations = []
                    for _ in range(row_count):
                        try:
                            location_tuple = fake.local_latlng(country_code=country_code, coords_only=coords_only)
                            if location_tuple:
                                locations.append(location_tuple)
                            else:
                                # Fallback if no location found
                                location_tuple = fake.local_latlng(country_code='US', coords_only=coords_only)
                                locations.append(location_tuple if location_tuple else ('40.712776', '-74.005974'))
                        except Exception:
                            # Ultimate fallback
                            if coords_only:
                                locations.append(('40.712776', '-74.005974'))
                            else:
                                locations.append(('40.712776', '-74.005974', 'New York', 'US', 'America/New_York'))
                    
                    # Split into separate columns
                    data[f"{col_name}_latitude"] = [str(loc[0]) for loc in locations]
                    data[f"{col_name}_longitude"] = [str(loc[1]) for loc in locations]
                    
                    if not coords_only and len(locations[0]) >= 5:
                        data[f"{col_name}_city"] = [str(loc[2]) for loc in locations]
                        data[f"{col_name}_country_code"] = [str(loc[3]) for loc in locations]
                        data[f"{col_name}_timezone"] = [str(loc[4]) for loc in locations]
                elif col_type == 'GenAI Text':
                    # For GenAI Text, initially fill with placeholder
                    # We'll update these after the DataFrame is created
                    prompt_template = col.get('prompt', '')
                    if prompt_template:
                        data[col_name] = [f"GenAI_Placeholder_{i}" for i in range(row_count)]
                    else:
                        data[col_name] = [""] * row_count
                elif col_type == 'Custom Values':
                    # For Custom Values, use values and optional weights with random selection
                    custom_values = col.get('custom_values', [''])
                    use_weights = col.get('use_weights', False)
                    custom_weights = col.get('custom_weights', [1])
                    
                    # Filter out empty values
                    filtered_values = [v for v in custom_values if v.strip()]
                    if not filtered_values:
                        filtered_values = ['DefaultValue']  # Fallback if all values are empty
                    
                    if use_weights and len(custom_weights) >= len(filtered_values):
                        # Use weights for random selection
                        filtered_weights = custom_weights[:len(filtered_values)]
                        # Normalize weights
                        total_weight = sum(filtered_weights)
                        if total_weight > 0:
                            probabilities = [w / total_weight for w in filtered_weights]
                        else:
                            probabilities = [1.0 / len(filtered_values)] * len(filtered_values)
                        
                        data[col_name] = [
                            random.choices(filtered_values, weights=probabilities)[0] 
                            for _ in range(row_count)
                        ]
                    else:
                        # Use uniform random selection without weights
                        data[col_name] = [random.choice(filtered_values) for _ in range(row_count)]
            
            # Only include user-specified columns (no automatic company columns)
            
            # Create DataFrame
            df = pd.DataFrame(data)
            
            # Process GenAI Text columns by calling the LLM endpoint
            from model_serving_utils import query_endpoint
            
            genai_columns = [col for col in columns if col.get('data_type') == 'GenAI Text']
            if genai_columns:
                print(f"🤖 PANDAS FALLBACK AI: Processing {len(genai_columns)} GenAI Text column(s)")
                print(f"   ✅ UNLIMITED: Will generate unique content for ALL {row_count} rows")
            
            for col in columns:
                col_name = col.get('name', 'unnamed_column')
                col_type = col.get('data_type', 'Integer')
                
                if col_type == 'GenAI Text':
                    prompt_template = col.get('prompt', '')
                    if prompt_template:
                        print(f"   - Column '{col_name}': Generating unique AI text for ALL {row_count} rows")
                        if row_count > 25:
                            print(f"   ⏱️  NOTE: This will make {row_count} individual API calls - may take a while")
                        self.generation_state['current_step'] = f"Generating AI text for column '{col_name}'..."
                        
                        # Generate text for each row - no longer limiting API calls
                        ai_texts = []
                        
                        for i in range(row_count):
                            try:
                                # Get current row data for column substitution
                                row_data = {col: df.iloc[i][col] for col in df.columns if col in df.columns}
                                
                                # Add table formatting note to the prompt
                                enhanced_prompt = f"{prompt_template} Note: This will be text data in a table so omit all special formatting."
                                
                                # Substitute column references with actual row values
                                contextualized_prompt = self._substitute_column_references_pandas(enhanced_prompt, row_data, columns)
                                
                                print(f"GenAI Text Debug - Row {i} contextualized prompt: {contextualized_prompt[:150]}...")
                                
                                # Use the LLM endpoint to generate text with higher token limit
                                messages = [{"role": "user", "content": contextualized_prompt}]
                                response = query_endpoint(self.endpoint_name, messages, 500)
                                
                                # Debug: log the raw response structure
                                print(f"GenAI Text Debug - Raw response type: {type(response)}")
                                print(f"GenAI Text Debug - Response preview: {str(response)[:200]}...")
                                
                                # Extract clean content from response using existing method
                                ai_text = self._extract_content_safely(response)
                                print(f"GenAI Text Debug - Extracted text: {ai_text[:100]}...")
                                
                                ai_texts.append(ai_text)
                            except Exception as e:
                                print(f"Error generating AI text for row {i}: {str(e)}")
                                ai_texts.append(f"Error generating text: {str(e)}")
                        
                        # Update the DataFrame column with all generated texts
                        df[col_name] = ai_texts
            
            # Save to CSV
            filename = f"{table_name}_{timestamp}.csv"
            local_dir = "./generated_documents"
            os.makedirs(local_dir, exist_ok=True)
            csv_path = os.path.join(local_dir, filename)
            
            df.to_csv(csv_path, index=False)
            
            # Get preview (first 5 rows)
            preview_data = df.head(5)
            
            # Save to volume
            volume_success = self._save_to_volume(csv_path, filename)
            
            return {
                'type': 'tabular',
                'table_name': table_name,
                'filename': filename,
                'path': csv_path,
                'row_count': row_count,
                'column_count': len(columns),
                'columns': [col.get('name', 'unnamed') for col in columns],
                'preview_data': preview_data.to_html(classes='table table-striped table-sm', index=False),
                'timestamp': timestamp,
                'company_name': company_name,
                'company_sector': company_sector,
                'size': os.path.getsize(csv_path),
                'saved_to_volume': volume_success
            }
            
        except Exception as e:
            print(f"Error in fallback tabular generation: {str(e)}")
            return {
                'type': 'tabular',
                'table_name': table_name,
                'filename': f"error_{timestamp}.csv",
                'error': str(e),
                'timestamp': timestamp,
                'company_name': company_name,
                'company_sector': company_sector,
                'saved_to_volume': False,
                'path': None,
                'size': 0
            }
